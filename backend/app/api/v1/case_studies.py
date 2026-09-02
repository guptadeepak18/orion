import io
import os
import re
import uuid
from typing import List, Optional, Dict, Any
from uuid import UUID
from fastapi import APIRouter, Depends, Query, status, HTTPException, UploadFile, File
from fastapi.responses import FileResponse, StreamingResponse, HTMLResponse
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import settings
from app.core.database import get_db
from app.core.permissions import require_permission, get_current_user, require_role
from app.models.auth import User
from app.schemas.common import ResponseEnvelope
from app.schemas.case_study import CaseStudyCreate, CaseStudyUpdate, CaseStudyResponse
from app.services.case_study_service import case_study_service

router = APIRouter(prefix="/case-studies", tags=["Case Studies Bank"])


@router.get(
    "",
    response_model=ResponseEnvelope[List[CaseStudyResponse]],
    dependencies=[Depends(require_permission("academic", "view_own"))],
)
async def list_case_studies(
    search: Optional[str] = Query(None, description="Search by title, concept, or domain"),
    domain: Optional[str] = Query(None, description="Filter by domain"),
    publisher: Optional[str] = Query(None, description="Filter by publisher/source"),
    tag: Optional[str] = Query(None, description="Filter by subject tag"),
    has_document: Optional[bool] = Query(None, description="Filter by document availability"),
    sort_by: Optional[str] = Query("newest", description="Sort by order (newest, oldest, title)"),
    db: AsyncSession = Depends(get_db),
):
    items = await case_study_service.list_case_studies(
        db,
        search=search,
        domain=domain,
        publisher=publisher,
        tag=tag,
        has_document=has_document,
        sort_by=sort_by,
        is_active=True,
    )
    return ResponseEnvelope(data=[CaseStudyResponse.model_validate(c) for c in items])


@router.get(
    "/{case_study_id}",
    response_model=ResponseEnvelope[CaseStudyResponse],
    dependencies=[Depends(require_permission("academic", "view_own"))],
)
async def get_case_study(
    case_study_id: str,
    db: AsyncSession = Depends(get_db),
):
    item = await case_study_service.get_case_study(db, case_study_id)
    if not item:
        raise HTTPException(status_code=404, detail="Case study not found")
    return ResponseEnvelope(data=CaseStudyResponse.model_validate(item))


@router.post(
    "",
    response_model=ResponseEnvelope[CaseStudyResponse],
    status_code=status.HTTP_201_CREATED,
    dependencies=[Depends(require_permission("academic", "edit"))],
)
async def create_case_study(
    payload: CaseStudyCreate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    cs = await case_study_service.create_case_study(db, payload, user_id=current_user.id)
    return ResponseEnvelope(data=CaseStudyResponse.model_validate(cs))


@router.put(
    "/{case_study_id}",
    response_model=ResponseEnvelope[CaseStudyResponse],
    dependencies=[Depends(require_permission("academic", "edit"))],
)
async def update_case_study(
    case_study_id: str,
    payload: CaseStudyUpdate,
    db: AsyncSession = Depends(get_db),
):
    cs = await case_study_service.update_case_study(db, case_study_id, payload)
    if not cs:
        raise HTTPException(status_code=404, detail="Case study not found")
    return ResponseEnvelope(data=CaseStudyResponse.model_validate(cs))


@router.delete(
    "/{case_study_id}",
    dependencies=[Depends(require_permission("academic", "edit"))],
)
async def delete_case_study(
    case_study_id: str,
    db: AsyncSession = Depends(get_db),
):
    success = await case_study_service.delete_case_study(db, case_study_id)
    if not success:
        raise HTTPException(status_code=404, detail="Case study not found")
    return ResponseEnvelope(data={"success": True, "message": "Case study archived"})


# ─── File Upload & Streaming Endpoints ────────────────────────────────────────

@router.post(
    "/upload-file",
    dependencies=[Depends(require_permission("academic", "edit"))],
)
async def upload_case_study_file(
    file: UploadFile = File(...),
):
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file uploaded")

    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="Empty file uploaded")

    upload_dir = os.path.join(settings.FILE_STORAGE_PATH, "case_studies")
    os.makedirs(upload_dir, exist_ok=True)

    file_id = uuid.uuid4().hex[:12]
    clean_name = re.sub(r'[^a-zA-Z0-9_\-\.]', '_', file.filename)
    stored_filename = f"{file_id}_{clean_name}"
    file_path = os.path.join(upload_dir, stored_filename)

    with open(file_path, "wb") as f:
        f.write(content)

    ext = os.path.splitext(file.filename)[1].lower().replace('.', '')

    return ResponseEnvelope(
        data={
            "file_url": f"/case-studies/files/{stored_filename}",
            "file_name": file.filename,
            "file_size": len(content),
            "file_type": ext,
            "stored_filename": stored_filename,
        }
    )


@router.get(
    "/files/{filename}",
    dependencies=[Depends(require_permission("academic", "view_own"))],
)
async def download_case_study_file(
    filename: str,
    original_name: Optional[str] = Query(None),
    inline: Optional[bool] = Query(False),
):
    upload_dir = os.path.join(settings.FILE_STORAGE_PATH, "case_studies")
    file_path = os.path.join(upload_dir, filename)

    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="Case study document not found on server")

    download_name = original_name or filename
    safe_name = re.sub(r'[^a-zA-Z0-9_\-\.]', '_', download_name)

    ext = os.path.splitext(filename)[1].lower().replace('.', '')
    media_types = {
        "pdf": "application/pdf",
        "txt": "text/plain",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "doc": "application/msword",
        "pptx": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
        "ppt": "application/vnd.ms-powerpoint",
        "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "zip": "application/zip",
    }
    m_type = media_types.get(ext, "application/octet-stream")
    disposition = "inline" if inline else f'attachment; filename="{safe_name}"'

    return FileResponse(
        path=file_path,
        media_type=m_type,
        filename=safe_name if not inline else None,
        headers={"Content-Disposition": disposition},
    )


# ─────────────────────────────────────────────────────────────────────────────
# AI CASE ANALYZER ENDPOINTS (PERSISTENT & USER-SPECIFIC)
# ─────────────────────────────────────────────────────────────────────────────
from app.services.case_analyzer_service import case_analyzer_service, FRAMEWORK_LENSES
from pydantic import BaseModel


class CaseAnalysisRequest(BaseModel):
    analysis_lens: str = "comprehensive_360"
    custom_prompt: Optional[str] = None


@router.get(
    "/meta/framework-lenses",
    response_model=ResponseEnvelope[List[Dict[str, Any]]],
    dependencies=[Depends(require_permission("academic", "view_own"))],
    summary="Get all available AI Case Analytical Framework Lenses",
)
async def get_framework_lenses():
    lenses = [
        {
            "id": k,
            "label": v["label"],
            "description": v["description"],
        }
        for k, v in FRAMEWORK_LENSES.items()
    ]
    return ResponseEnvelope(data=lenses)


@router.get(
    "/{case_study_id}/ai-analysis",
    response_model=ResponseEnvelope[Optional[Dict[str, Any]]],
    dependencies=[Depends(require_permission("academic", "view_own"))],
    summary="Get user-specific saved AI Case Analysis",
)
async def get_case_analysis(
    case_study_id: UUID,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    analysis = await case_analyzer_service.get_user_analysis(
        db=db,
        case_study_id=case_study_id,
        user_id=current_user.id,
    )
    return ResponseEnvelope(data=analysis)


@router.post(
    "/{case_study_id}/ai-analysis",
    response_model=ResponseEnvelope[Dict[str, Any]],
    dependencies=[Depends(require_permission("academic", "view_own"))],
    summary="Run/Re-run AI Case Analyzer and save user-specific results",
)
async def run_case_analysis(
    case_study_id: UUID,
    payload: CaseAnalysisRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    try:
        result = await case_analyzer_service.generate_and_save_analysis(
            db=db,
            case_study_id=case_study_id,
            user_id=current_user.id,
            analysis_lens=payload.analysis_lens,
            custom_prompt=payload.custom_prompt,
        )
        return ResponseEnvelope(data=result)
    except ValueError as e:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=str(e))


@router.delete(
    "/{case_study_id}/ai-analysis",
    dependencies=[Depends(require_permission("academic", "view_own"))],
    summary="Delete user-specific saved AI Case Analysis",
)
async def delete_case_analysis(
    case_study_id: UUID,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    success = await case_analyzer_service.delete_user_analysis(
        db=db,
        case_study_id=case_study_id,
        user_id=current_user.id,
    )
    return ResponseEnvelope(data={"deleted": success})


@router.get(
    "/{case_study_id}/ai-analysis/export",
    dependencies=[Depends(require_permission("academic", "view_own"))],
    summary="Export AI Case Analysis as DOCX or print-ready HTML",
)
async def export_case_analysis(
    case_study_id: UUID,
    format: str = Query("docx", description="Export format: 'docx' or 'html'"),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    analysis_data = await case_analyzer_service.get_user_analysis(
        db=db,
        case_study_id=case_study_id,
        user_id=current_user.id,
    )
    if not analysis_data:
        raise HTTPException(status_code=404, detail="No saved analysis found for this case study. Please run the AI analyzer first.")

    analysis = analysis_data.get("analysis", {})
    meta = analysis.get("meta", {})
    sections = analysis.get("sections", [])
    case_title = meta.get("case_title", "Case Analysis")
    framework_lens = meta.get("framework_lens", "")
    generated_timestamp = meta.get("generated_timestamp", "")
    safe_title = re.sub(r'[^\w\s-]', '', case_title).strip().replace(' ', '_')[:60]

    if format == "docx":
        return _generate_docx_response(analysis, case_title, framework_lens, generated_timestamp, safe_title)
    else:
        user_name = current_user.full_name or current_user.email or "Faculty Member, Lexicon MILE"
        return _generate_html_response(analysis, case_title, framework_lens, generated_timestamp, user_name)


def _generate_docx_response(analysis: dict, case_title: str, framework_lens: str, generated_timestamp: str, safe_title: str) -> StreamingResponse:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx not installed. Run: pip install python-docx")

    doc = Document()

    # --- Page margins ---
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # --- Colors ---
    NAVY = RGBColor(0x0F, 0x29, 0x4A)      # Primary Title
    BLUE = RGBColor(0x1D, 0x4E, 0xD8)      # Heading 1
    EMERALD = RGBColor(0x04, 0x78, 0x57)   # Outcomes
    SLATE = RGBColor(0x47, 0x55, 0x69)     # Subtext
    DARK = RGBColor(0x0F, 0x17, 0x2A)      # Body text

    # --- Cover Title Block ---
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("CASE STUDY STRATEGIC MASTERCLASS & OUTCOME BRIEF")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = NAVY

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle_para.add_run(case_title)
    sr.font.size = Pt(15)
    sr.font.bold = True
    sr.font.color.rgb = DARK

    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta_para.add_run(f"Framework: {framework_lens}   |   Generated: {generated_timestamp}   |   MBA / Executive Tier")
    mr.font.size = Pt(9)
    mr.font.color.rgb = SLATE
    mr.font.italic = True

    doc.add_paragraph().add_run("═" * 70).font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    doc.add_paragraph()

    # --- Check format: New Object-Based vs Legacy Sections Array ---
    snap = analysis.get("executive_snapshot", {})
    conflict = analysis.get("conflict_and_market_forces", {})
    alts = analysis.get("strategic_alternatives", {})
    outcome = analysis.get("real_world_outcome", {})
    root = analysis.get("root_cause_diagnostic", {})
    takeaways = analysis.get("key_business_takeaways", {})
    exam = analysis.get("exam_and_viva_prep", {})

    if snap or outcome:
        # ── 1. EXECUTIVE SNAPSHOT ──
        h1 = doc.add_heading("1. Executive Snapshot & The Burning Platform", level=1)
        h1.runs[0].font.color.rgb = BLUE
        if snap.get("headline"):
            hp = doc.add_paragraph()
            hr = hp.add_run(f"Executive Summary: {snap['headline']}")
            hr.font.bold = True
            hr.font.size = Pt(10.5)
            hr.font.color.rgb = NAVY
        if snap.get("burning_platform"):
            doc.add_paragraph(snap["burning_platform"]).runs[0].font.size = Pt(10)
        if snap.get("core_dilemma"):
            dp = doc.add_paragraph()
            dp.add_run("Core Strategic Dilemma: ").bold = True
            dp.add_run(snap["core_dilemma"]).font.size = Pt(10)

        if snap.get("key_actors"):
            doc.add_heading("Key Decision Makers & Protagonists", level=2)
            for actor in snap["key_actors"]:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"{actor.get('actor', '')}: ").bold = True
                p.add_run(actor.get("position", ""))

        doc.add_paragraph()

        # ── 2. CONFLICT & MARKET FORCES ──
        h2 = doc.add_heading("2. Conflict, Market Forces & Ecosystem Breakdown", level=1)
        h2.runs[0].font.color.rgb = BLUE
        if conflict.get("narrative"):
            doc.add_paragraph(conflict["narrative"]).runs[0].font.size = Pt(10)
        if conflict.get("key_tensions"):
            doc.add_heading("Key Strategic Tensions", level=2)
            for t in conflict["key_tensions"]:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"{t.get('tension_title', '')}: ").bold = True
                p.add_run(t.get("description", ""))
        if conflict.get("stakeholder_matrix"):
            doc.add_heading("Stakeholder Impact Matrix", level=2)
            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Stakeholder", "Leverage", "Core Priority", "Outcome"
            for st in conflict["stakeholder_matrix"]:
                row = table.add_row().cells
                row[0].text = st.get("stakeholder", "")
                row[1].text = st.get("leverage", "")
                row[2].text = st.get("priority", "")
                row[3].text = st.get("outcome", "")

        doc.add_paragraph()

        # ── 3. STRATEGIC ALTERNATIVES MATRIX ──
        h3 = doc.add_heading("3. Strategic Alternatives & Decision Matrix", level=1)
        h3.runs[0].font.color.rgb = BLUE
        if alts.get("framing"):
            doc.add_paragraph(alts["framing"]).runs[0].font.size = Pt(10)
        if alts.get("options"):
            table = doc.add_table(rows=1, cols=5)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text, hdr[4].text = "Strategy Option", "Key Advantages (Pros)", "Risks & Downsides (Cons)", "Risk Tier", "Capital & Time"
            for opt in alts["options"]:
                row = table.add_row().cells
                row[0].text = opt.get("option_name", "")
                row[1].text = "\n".join([f"• {p}" for p in opt.get("pros", [])]) if isinstance(opt.get("pros"), list) else str(opt.get("pros", ""))
                row[2].text = "\n".join([f"• {c}" for c in opt.get("cons", [])]) if isinstance(opt.get("cons"), list) else str(opt.get("cons", ""))
                row[3].text = opt.get("risk_level", "")
                row[4].text = opt.get("capital_and_time", "")

        doc.add_paragraph()

        # ── 4. THE REAL-WORLD OUTCOME (CRITICAL) ──
        h4 = doc.add_heading("4. The Real-World Outcome & Historical Resolution", level=1)
        h4.runs[0].font.color.rgb = EMERALD
        if outcome.get("actual_decision"):
            dp = doc.add_paragraph()
            r1 = dp.add_run("The Actual Strategic Path Chosen: ")
            r1.bold = True
            r1.font.color.rgb = EMERALD
            dp.add_run(outcome["actual_decision"]).font.size = Pt(10.5)

        if outcome.get("execution_and_aftermath"):
            doc.add_heading("Execution Trajectory & Market Aftermath", level=2)
            doc.add_paragraph(outcome["execution_and_aftermath"]).runs[0].font.size = Pt(10)

        if outcome.get("measurable_results"):
            doc.add_heading("Quantifiable Outcomes & Business Metrics", level=2)
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text = "Outcome Metric / Event", "Measured Value / Terms", "Strategic Significance"
            for m in outcome["measurable_results"]:
                row = table.add_row().cells
                row[0].text = m.get("metric", "")
                row[1].text = m.get("value", "")
                row[2].text = m.get("significance", "")

        if outcome.get("long_term_legacy"):
            doc.add_heading("Enduring Legacy & Industry Impact", level=2)
            doc.add_paragraph(outcome["long_term_legacy"]).runs[0].font.size = Pt(10)

        doc.add_paragraph()

        # ── 5. ROOT CAUSE DIAGNOSTIC ──
        if root.get("why_it_happened") or root.get("critical_failure_or_success_factors"):
            h5 = doc.add_heading("5. Root-Cause Diagnostic: Why It Succeeded or Failed", level=1)
            h5.runs[0].font.color.rgb = BLUE
            if root.get("why_it_happened"):
                doc.add_paragraph(root["why_it_happened"]).runs[0].font.size = Pt(10)
            if root.get("critical_failure_or_success_factors"):
                doc.add_heading("Critical Success / Failure Determinants", level=2)
                for f in root["critical_failure_or_success_factors"]:
                    doc.add_paragraph(f, style="List Bullet")
            doc.add_paragraph()

        # ── 6. KEY BUSINESS TAKEAWAYS & MENTAL MODELS ──
        h6 = doc.add_heading("6. Key Business Principles & Strategic Mental Models", level=1)
        h6.runs[0].font.color.rgb = NAVY
        if takeaways.get("summary"):
            doc.add_paragraph(takeaways["summary"]).runs[0].font.size = Pt(10)
        if takeaways.get("principles"):
            for p in takeaways["principles"]:
                doc.add_heading(p.get("principle_name", "Principle"), level=2)
                if p.get("concept"):
                    p_para = doc.add_paragraph()
                    p_para.add_run("Core Concept: ").bold = True
                    p_para.add_run(p["concept"])
                if p.get("application"):
                    a_para = doc.add_paragraph()
                    a_para.add_run("Managerial Application: ").bold = True
                    a_para.add_run(p["application"])

        doc.add_paragraph()

        # ── 7. EXAM & SOCRATIC VIVA PREP ──
        h7 = doc.add_heading("7. Classroom Socratic Discussion & Exam Prep Guide", level=1)
        h7.runs[0].font.color.rgb = BLUE
        if exam.get("facilitation_note"):
            fp = doc.add_paragraph()
            fp.add_run("Facilitator & Student Guidance: ").bold = True
            fp.add_run(exam["facilitation_note"]).italic = True
        if exam.get("questions_and_answers"):
            for i, qa in enumerate(exam["questions_and_answers"], 1):
                qp = doc.add_paragraph()
                qp.add_run(f"Q{i}: {qa.get('question', '')}").bold = True
                if qa.get("model_answer"):
                    ap = doc.add_paragraph()
                    ap.add_run("Model Student Answer: ").bold = True
                    ap.add_run(qa["model_answer"])
                if qa.get("framework_to_cite"):
                    fp = doc.add_paragraph()
                    fp.add_run(f"Framework to Cite: {qa['framework_to_cite']}").font.color.rgb = SLATE
                    fp.runs[0].font.italic = True
                doc.add_paragraph()

    else:
        # Fallback for old 8-section array format
        sections = analysis.get("sections", [])
        for sec in sections:
            doc.add_heading(sec.get("title", ""), level=1)
            for p in sec.get("paragraphs", []):
                doc.add_paragraph(p)
            for p in sec.get("summary_points", []):
                doc.add_paragraph(p, style="List Bullet")
            for t in sec.get("takeaways", []):
                doc.add_paragraph(t, style="List Bullet")
            for q in sec.get("questions", []):
                doc.add_paragraph(q, style="List Number")

    # --- Footer note ---
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_para.add_run(f"Generated by Orion AI Strategic Case Intelligence Engine  •  {generated_timestamp}")
    fr.font.size = Pt(8)
    fr.font.color.rgb = SLATE
    fr.font.italic = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="AI_Case_Analysis_{safe_title}.docx"'},
    )


def _generate_html_response(
    analysis: dict,
    case_title: str,
    framework_lens: str,
    generated_timestamp: str,
    authorized_for: str = "Faculty Member, Lexicon MILE",
) -> HTMLResponse:
    import html

    def esc(s: Any) -> str:
        return html.escape(str(s or ""))

    meta = analysis.get("meta", {})
    snap = analysis.get("executive_snapshot", {})
    conflict = analysis.get("conflict_and_market_forces", {})
    alts = analysis.get("strategic_alternatives", {})
    outcome = analysis.get("real_world_outcome", {})
    root = analysis.get("root_cause_diagnostic", {})
    takeaways = analysis.get("key_business_takeaways", {})
    exam = analysis.get("exam_and_viva_prep", {})

    case_name = esc(meta.get("case_title", case_title))
    auth_name = esc(authorized_for)
    gen_time = esc(meta.get("generated_timestamp", generated_timestamp))
    lens_label = esc(meta.get("framework_lens", framework_lens or "Outcome Mastery"))
    case_code = esc(meta.get("case_code", "HBP-STRAT"))
    setting = esc(meta.get("setting", "Global Strategy & Leadership"))

    html_parts = [f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>AI Case Masterclass — {case_name}</title>
<style>
  @page {{
    size: letter;
    margin: 1.8cm 1.8cm 1.8cm 1.8cm;
  }}
  body {{
    font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", Arial, sans-serif;
    color: #1e293b;
    line-height: 1.6;
    background: #f8fafc;
    margin: 0;
    padding: 24px;
    font-size: 10pt;
  }}
  .document-container {{
    max-width: 860px;
    margin: 0 auto;
    background: #ffffff;
    padding: 48px 56px;
    border: 1px solid #e2e8f0;
    box-shadow: 0 4px 20px rgba(0,0,0,0.06);
    border-radius: 4px;
  }}
  .hbp-top-banner {{
    border-bottom: 2px solid #0f172a;
    padding-bottom: 12px;
    margin-bottom: 24px;
  }}
  .hbp-dnc {{
    font-size: 9pt;
    font-weight: 800;
    text-transform: uppercase;
    letter-spacing: 0.05em;
    color: #b91c1c;
    margin-bottom: 2px;
  }}
  .hbp-created-for {{
    font-size: 8.5pt;
    color: #64748b;
    font-style: italic;
  }}
  .hbp-main-title {{
    font-family: "Georgia", serif;
    font-size: 22pt;
    font-weight: 900;
    color: #0f294a;
    margin: 16px 0 8px 0;
    letter-spacing: -0.01em;
    line-height: 1.2;
  }}
  .hbp-supporting {{
    font-size: 11pt;
    font-weight: 700;
    color: #334155;
    margin-bottom: 14px;
  }}
  .hbp-supporting-sub {{
    font-size: 10pt;
    font-weight: normal;
    color: #475569;
  }}
  .hbp-auth-box {{
    background: #f1f5f9;
    border-left: 3px solid #0f294a;
    padding: 10px 14px;
    font-size: 8.5pt;
    color: #475569;
    margin: 14px 0 20px 0;
    line-height: 1.5;
  }}
  .hbp-meta-bar {{
    font-size: 9pt;
    font-weight: 600;
    color: #475569;
    padding: 8px 0 0 0;
    border-top: 1px solid #e2e8f0;
  }}
  .sec-h1 {{
    font-family: "Georgia", serif;
    font-size: 13pt;
    font-weight: 800;
    color: #0f294a;
    margin: 28px 0 10px 0;
    padding-bottom: 4px;
    border-bottom: 1.5px solid #0f294a;
    letter-spacing: -0.01em;
    page-break-after: avoid;
  }}
  .sec-h1.outcome-h1 {{
    color: #065f46;
    border-bottom-color: #065f46;
  }}
  .sec-h2 {{
    font-size: 10.5pt;
    font-weight: 700;
    color: #1e3a8a;
    margin: 16px 0 6px 0;
    page-break-after: avoid;
  }}
  p {{
    margin: 0 0 10px 0;
    text-align: justify;
    line-height: 1.6;
  }}
  .highlight-card {{
    background: #eff6ff;
    border-left: 4px solid #2563eb;
    padding: 12px 16px;
    margin: 12px 0;
    font-weight: 600;
    color: #1e3a8a;
    border-radius: 2px;
    line-height: 1.5;
  }}
  .decision-box {{
    background: #f8fafc;
    border: 1px solid #cbd5e1;
    border-left: 4px solid #d97706;
    padding: 12px 16px;
    margin: 12px 0;
    font-weight: 600;
    color: #92400e;
    border-radius: 2px;
    line-height: 1.5;
  }}
  .outcome-card {{
    background: #f0fdf4;
    border: 1px solid #bbf7d0;
    border-left: 5px solid #16a34a;
    padding: 16px 20px;
    margin: 14px 0;
    border-radius: 4px;
  }}
  .outcome-card h3 {{
    font-size: 11pt;
    font-weight: 800;
    color: #15803d;
    margin: 0 0 6px 0;
  }}
  table {{
    width: 100%;
    border-collapse: collapse;
    margin: 14px 0;
    font-size: 9pt;
    page-break-inside: avoid;
  }}
  th {{
    background: #0f294a;
    color: #ffffff;
    font-weight: 700;
    text-align: left;
    padding: 8px 10px;
    border: 1px solid #0f294a;
  }}
  th.outcome-th {{
    background: #065f46;
    border-color: #065f46;
  }}
  td {{
    padding: 7px 10px;
    border: 1px solid #e2e8f0;
    vertical-align: top;
    line-height: 1.45;
  }}
  tr:nth-child(even) td {{
    background: #f8fafc;
  }}
  .takeaway-card {{
    background: #fdf4ff;
    border-left: 4px solid #a855f7;
    border: 1px solid #f3e8ff;
    border-left-width: 4px;
    border-left-color: #9333ea;
    border-radius: 3px;
    padding: 12px 16px;
    margin: 10px 0;
    page-break-inside: avoid;
  }}
  .takeaway-title {{
    font-size: 10pt;
    font-weight: 800;
    color: #6b21a8;
    margin-bottom: 4px;
  }}
  .takeaway-body {{
    font-size: 9pt;
    color: #4c1d95;
    line-height: 1.5;
  }}
  .q-block {{
    background: #ffffff;
    border: 1px solid #e2e8f0;
    border-left: 3px solid #0f294a;
    padding: 14px 16px;
    margin: 14px 0;
    page-break-inside: avoid;
  }}
  .q-num {{
    font-weight: 800;
    color: #0f294a;
    font-size: 10.5pt;
    margin-bottom: 4px;
  }}
  .q-text {{
    font-weight: 700;
    font-style: italic;
    color: #0f172a;
    margin-bottom: 8px;
    font-size: 10pt;
  }}
  .q-ans {{
    font-size: 9pt;
    color: #334155;
    line-height: 1.55;
    margin-bottom: 6px;
  }}
  .q-framework {{
    font-size: 8.5pt;
    color: #2563eb;
    font-weight: 600;
  }}
  .footer-bar {{
    margin-top: 36px;
    padding-top: 12px;
    border-top: 1px solid #cbd5e1;
    font-size: 8pt;
    color: #64748b;
    display: flex;
    justify-content: space-between;
  }}
  @media print {{
    body {{
      background: #ffffff;
      padding: 0;
      font-size: 9.5pt;
    }}
    .document-container {{
      border: none;
      box-shadow: none;
      padding: 0;
      max-width: 100%;
    }}
    .no-print {{
      display: none !important;
    }}
    .sec-h1, .sec-h2, .q-block, .outcome-card, .takeaway-card {{
      page-break-inside: avoid;
    }}
    th, th.outcome-th {{
      -webkit-print-color-adjust: exact;
      print-color-adjust: exact;
    }}
    .highlight-card, .decision-box, .outcome-card {{
      -webkit-print-color-adjust: exact;
      print-color-adjust: exact;
    }}
  }}
</style>
</head>
<body>

<div class="no-print" style="max-width:860px;margin:0 auto 16px auto;display:flex;justify-content:space-between;align-items:center;padding:10px 16px;background:#0f294a;color:#fff;border-radius:6px;">
  <div><strong>AI Case Study Masterclass View</strong> — Authorized Executive Brief</div>
  <button onclick="window.print()" style="background:#2563eb;color:#fff;border:none;padding:6px 14px;border-radius:4px;font-weight:bold;cursor:pointer;">🖨️ Print / Save as PDF</button>
</div>

<div class="document-container">
  <div class="hbp-top-banner">
    <div class="hbp-dnc">Do Not Copy or Post</div>
    <div class="hbp-created-for">Case Masterclass created for {auth_name} on {gen_time}</div>
    <div class="hbp-main-title">AI Case Masterclass — {case_name}</div>
    <div class="hbp-supporting">SUPPORTING:<br><span class="hbp-supporting-sub">{case_name}<br>Analytical Framework Lens: {lens_label}</span></div>
    
    <div class="hbp-auth-box">
      This Case Study Outcome Masterclass is authorized for <strong>{auth_name}, Lexicon MILE - Management Institute of Leadership and Excellence</strong>.<br>
      <span style="font-size:8pt;color:#64748b;">This analysis was created by generative AI. Although we strive for accuracy, its responses may be inaccurate or incomplete. Please review all information provided carefully.</span>
    </div>

    <div class="hbp-meta-bar">
      Case: <strong>{case_code}</strong> &nbsp;|&nbsp; Lens: <strong>{lens_label}</strong> &nbsp;|&nbsp; Setting: <strong>{setting}</strong> &nbsp;|&nbsp; Tier: <strong>Executive / MBA Tier</strong>
    </div>
  </div>
"""]

    # 1. EXECUTIVE SNAPSHOT
    if snap:
        html_parts.append('<h1 class="sec-h1">1. Executive Snapshot & The Burning Platform</h1>')
        if snap.get("headline"):
            html_parts.append(f'<div class="highlight-card">{esc(snap["headline"])}</div>')
        if snap.get("burning_platform"):
            html_parts.append(f'<p>{esc(snap["burning_platform"])}</p>')
        if snap.get("core_dilemma"):
            html_parts.append(f'<div class="decision-box"><strong>Core Strategic Dilemma:</strong> {esc(snap["core_dilemma"])}</div>')
        if snap.get("key_actors"):
            html_parts.append('<h2 class="sec-h2">Key Decision Makers & Protagonists</h2>')
            html_parts.append('<table><tr><th style="width:30%;">Decision Maker</th><th>Position & Strategic Perspective</th></tr>')
            for a in snap["key_actors"]:
                html_parts.append(f'<tr><td><strong>{esc(a.get("actor",""))}</strong></td><td>{esc(a.get("position",""))}</td></tr>')
            html_parts.append('</table>')

    # 2. CONFLICT & MARKET FORCES
    if conflict:
        html_parts.append('<h1 class="sec-h1">2. Conflict, Market Forces & Ecosystem Breakdown</h1>')
        if conflict.get("narrative"):
            html_parts.append(f'<p>{esc(conflict["narrative"])}</p>')
        if conflict.get("key_tensions"):
            html_parts.append('<h2 class="sec-h2">Key Strategic Tensions</h2>')
            html_parts.append('<table><tr><th style="width:30%;">Tension / Dynamic</th><th>Description & Systemic Impact</th></tr>')
            for t in conflict["key_tensions"]:
                html_parts.append(f'<tr><td><strong>{esc(t.get("tension_title",""))}</strong></td><td>{esc(t.get("description",""))}</td></tr>')
            html_parts.append('</table>')
        if conflict.get("stakeholder_matrix"):
            html_parts.append('<h2 class="sec-h2">Stakeholder Impact Matrix</h2>')
            html_parts.append('<table><tr><th>Stakeholder</th><th>Leverage</th><th>Core Priority</th><th>Outcome</th></tr>')
            for s in conflict["stakeholder_matrix"]:
                html_parts.append(f'<tr><td><strong>{esc(s.get("stakeholder",""))}</strong></td><td>{esc(s.get("leverage",""))}</td><td>{esc(s.get("priority",""))}</td><td>{esc(s.get("outcome",""))}</td></tr>')
            html_parts.append('</table>')

    # 3. STRATEGIC ALTERNATIVES
    if alts:
        html_parts.append('<h1 class="sec-h1">3. Strategic Alternatives & Decision Matrix</h1>')
        if alts.get("framing"):
            html_parts.append(f'<p>{esc(alts["framing"])}</p>')
        if alts.get("options"):
            html_parts.append('<table><tr><th style="width:25%;">Strategic Option</th><th>Advantages (Pros)</th><th>Risks (Cons)</th><th style="width:12%;">Risk Tier</th><th style="width:18%;">Capital & Timeline</th></tr>')
            for o in alts["options"]:
                pros = "<br>• ".join([esc(p) for p in o.get("pros", [])]) if isinstance(o.get("pros"), list) else esc(o.get("pros", ""))
                cons = "<br>• ".join([esc(c) for c in o.get("cons", [])]) if isinstance(o.get("cons"), list) else esc(o.get("cons", ""))
                html_parts.append(f'<tr><td><strong>{esc(o.get("option_name",""))}</strong><br><small style="color:#64748b;">{esc(o.get("description",""))}</small></td><td>• {pros}</td><td>• {cons}</td><td><strong>{esc(o.get("risk_level",""))}</strong></td><td>{esc(o.get("capital_and_time",""))}</td></tr>')
            html_parts.append('</table>')

    # 4. THE REAL-WORLD OUTCOME (STAR SECTION)
    if outcome:
        html_parts.append('<h1 class="sec-h1 outcome-h1">4. The Real-World Outcome & Historical Resolution</h1>')
        if outcome.get("actual_decision"):
            html_parts.append(f'<div class="outcome-card"><h3>The Actual Strategic Path Chosen:</h3><p style="font-size:10pt;font-weight:600;color:#14532d;margin:0;">{esc(outcome["actual_decision"])}</p></div>')
        if outcome.get("execution_and_aftermath"):
            html_parts.append(f'<p>{esc(outcome["execution_and_aftermath"])}</p>')
        if outcome.get("measurable_results"):
            html_parts.append('<h2 class="sec-h2" style="color:#065f46;">Measurable Results & Hard Metrics</h2>')
            html_parts.append('<table><tr><th class="outcome-th">Outcome Metric / Event</th><th class="outcome-th">Measured Value / Terms</th><th class="outcome-th">Strategic Significance</th></tr>')
            for m in outcome["measurable_results"]:
                html_parts.append(f'<tr><td><strong>{esc(m.get("metric",""))}</strong></td><td><span style="color:#047857;font-weight:700;">{esc(m.get("value",""))}</span></td><td>{esc(m.get("significance",""))}</td></tr>')
            html_parts.append('</table>')
        if outcome.get("long_term_legacy"):
            html_parts.append(f'<p><strong>Long-Term Industry Legacy:</strong> {esc(outcome["long_term_legacy"])}</p>')

    # 5. ROOT CAUSE DIAGNOSTIC
    if root:
        html_parts.append('<h1 class="sec-h1">5. Root-Cause Diagnostic: Why It Succeeded or Failed</h1>')
        if root.get("why_it_happened"):
            html_parts.append(f'<p>{esc(root["why_it_happened"])}</p>')
        if root.get("critical_failure_or_success_factors"):
            html_parts.append('<h2 class="sec-h2">Critical Success & Failure Factors</h2><ul>')
            for f in root["critical_failure_or_success_factors"]:
                html_parts.append(f'<li>{esc(f)}</li>')
            html_parts.append('</ul>')

    # 6. KEY BUSINESS PRINCIPLES
    if takeaways:
        html_parts.append('<h1 class="sec-h1">6. Key Business Principles & Strategic Mental Models</h1>')
        if takeaways.get("summary"):
            html_parts.append(f'<p>{esc(takeaways["summary"])}</p>')
        if takeaways.get("principles"):
            for p in takeaways["principles"]:
                html_parts.append(f'<div class="takeaway-card"><div class="takeaway-title">{esc(p.get("principle_name",""))}</div><div class="takeaway-body"><strong>Core Concept:</strong> {esc(p.get("concept",""))}<br><br><strong>Managerial Application:</strong> {esc(p.get("application",""))}</div></div>')

    # 7. EXAM & SOCRATIC DISCUSSION PREP
    if exam:
        html_parts.append('<h1 class="sec-h1">7. Classroom Socratic Discussion & Exam Prep Guide</h1>')
        if exam.get("facilitation_note"):
            html_parts.append(f'<p><em><strong>Facilitator Note:</strong> {esc(exam["facilitation_note"])}</em></p>')
        if exam.get("questions_and_answers"):
            for i, qa in enumerate(exam["questions_and_answers"], 1):
                html_parts.append(f'<div class="q-block"><div class="q-num">Question {i}</div><div class="q-text">{esc(qa.get("question",""))}</div><div class="q-ans"><strong>Model Student Answer:</strong> {esc(qa.get("model_answer",""))}</div><div class="q-framework">Framework to Cite: {esc(qa.get("framework_to_cite",""))}</div></div>')

    html_parts.append(f"""
    <div class="footer-bar">
      <div>AI Case Study Masterclass — {case_name}</div>
      <div>Authorized for {auth_name}</div>
    </div>
  </div>

  <script>
    window.addEventListener('load', function() {{
      setTimeout(function() {{
        try {{
          window.print();
        }} catch(e) {{
          console.error('Print auto-trigger notice:', e);
        }}
      }}, 700);
    }});
  </script>
</body>
</html>""")

    return HTMLResponse(content="".join(html_parts))


# ─────────────────────────────────────────────────────────────────────────────
# HARVARD AI CASE NOTE GENERATOR ENDPOINTS (HBP TEACHING NOTE SPECIFICATION)
# ─────────────────────────────────────────────────────────────────────────────
from app.services.case_note_service import case_note_service


class CaseNoteGenerateRequest(BaseModel):
    duration_minutes: int = 80
    course_focus: str = "Strategy & General Management"
    custom_prompt: Optional[str] = None


@router.get(
    "/{case_study_id}/ai-case-note",
    response_model=ResponseEnvelope[Optional[Dict[str, Any]]],
    dependencies=[Depends(require_role(["super_admin", "admin", "crc_admin", "crc_coordinator", "faculty_internal", "faculty_external", "director"]))],
    summary="Get user-specific saved Harvard AI Case Note (Faculty & Admin only)",
)
async def get_case_note(
    case_study_id: UUID,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    note = await case_note_service.get_user_case_note(
        db=db,
        case_study_id=case_study_id,
        user_id=current_user.id,
    )
    return ResponseEnvelope(data=note)


@router.post(
    "/{case_study_id}/ai-case-note",
    response_model=ResponseEnvelope[Dict[str, Any]],
    dependencies=[Depends(require_role(["super_admin", "admin", "crc_admin", "crc_coordinator", "faculty_internal", "faculty_external", "director"]))],
    summary="Generate or customize Harvard AI Case Note with preferences (Faculty & Admin only)",
)
async def generate_case_note(
    case_study_id: UUID,
    payload: CaseNoteGenerateRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    try:
        user_name = current_user.full_name or current_user.email or "Faculty / Instructor"
        result = await case_note_service.generate_and_save_case_note(
            db=db,
            case_study_id=case_study_id,
            user_id=current_user.id,
            duration_minutes=payload.duration_minutes,
            course_focus=payload.course_focus,
            custom_prompt=payload.custom_prompt,
            authorized_user_name=user_name,
        )
        return ResponseEnvelope(data=result)
    except ValueError as e:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=str(e))


@router.get(
    "/{case_study_id}/ai-case-note/export",
    dependencies=[Depends(require_role(["super_admin", "admin", "crc_admin", "crc_coordinator", "faculty_internal", "faculty_external", "director"]))],
    summary="Export Harvard AI Case Note as DOCX or Print-Ready PDF / HTML (Faculty & Admin only)",
)
async def export_case_note(
    case_study_id: UUID,
    format: str = Query("docx", description="Export format: 'docx', 'html', or 'pdf'"),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    note_record = await case_note_service.get_user_case_note(
        db=db,
        case_study_id=case_study_id,
        user_id=current_user.id,
    )
    if not note_record:
        user_name = current_user.full_name or current_user.email or "Faculty / Student"
        note_record = await case_note_service.generate_and_save_case_note(
            db=db,
            case_study_id=case_study_id,
            user_id=current_user.id,
            authorized_user_name=user_name,
        )

    case_note = note_record.get("case_note", {})
    meta = case_note.get("meta", {})
    case_title = meta.get("case_title", "Harvard AI Case Note")
    safe_title = re.sub(r'[^\w\s-]', '', case_title).strip().replace(' ', '_')[:60]

    if format == "docx":
        return _generate_case_note_docx_response(case_note, case_title, safe_title)
    else:
        return _generate_case_note_html_response(case_note, case_title)


def _generate_case_note_docx_response(case_note: dict, case_title: str, safe_title: str) -> StreamingResponse:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx library not installed on server.")

    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    meta = case_note.get("meta", {})
    overview = case_note.get("overview", {})
    learning_objs = case_note.get("learning_objectives", [])
    why_teach = case_note.get("why_teach_this_case", {})
    theories = case_note.get("theory_and_frameworks", [])
    strategy = case_note.get("classroom_strategy", {})
    disc_plan = case_note.get("discussion_plan", {})
    takeaways = case_note.get("takeaways_and_application", {})
    alternatives = case_note.get("alternative_approaches", [])

    auth_name = meta.get("authorized_for", "Faculty Member, Lexicon MILE")

    hdr_p = doc.add_paragraph()
    r1 = hdr_p.add_run("Do Not Copy or Post\n")
    r1.font.size = Pt(9)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(160, 30, 30)

    r2 = hdr_p.add_run(f"Case Note created for {auth_name} on {meta.get('generated_timestamp', '')}\n")
    r2.font.size = Pt(8.5)
    r2.font.italic = True
    r2.font.color.rgb = RGBColor(100, 100, 100)

    title_p = doc.add_paragraph()
    tr = title_p.add_run(f"AI Case Note — {meta.get('case_title', case_title)}")
    tr.font.size = Pt(18)
    tr.font.bold = True
    tr.font.color.rgb = RGBColor(15, 41, 74)

    sup_p = doc.add_paragraph()
    sr = sup_p.add_run(f"SUPPORTING:\n{meta.get('case_title', case_title)}\nBy {meta.get('authors', 'Case Authors')}")
    sr.font.size = Pt(10.5)
    sr.font.bold = True
    sr.font.color.rgb = RGBColor(70, 70, 70)

    disc_p = doc.add_paragraph()
    disc_r = disc_p.add_run(f"This Case Note is authorized for {auth_name}.\n"
                            "This case note was created by generative AI. Although we strive for accuracy, its responses may be inaccurate or incomplete. Please review all information provided carefully.\n"
                            f"Case: {meta.get('case_code', 'HBP-CASE')} | Authors: {meta.get('authors', '')} | Setting: {meta.get('setting', '')}")
    disc_r.font.size = Pt(8)
    disc_r.font.color.rgb = RGBColor(110, 110, 110)

    doc.add_paragraph("―" * 45)

    def add_h1(text: str):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(4)
        r = h.add_run(text)
        r.font.size = Pt(13)
        r.font.bold = True
        r.font.color.rgb = RGBColor(15, 41, 74)

    def add_h2(text: str):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(3)
        r = h.add_run(text)
        r.font.size = Pt(11)
        r.font.bold = True
        r.font.color.rgb = RGBColor(40, 60, 100)

    # 1. CASE OVERVIEW
    if overview:
        add_h1("1. Case Overview")
        if overview.get("story_synopsis"):
            add_h2("Story Synopsis")
            for para in overview["story_synopsis"].split("\n\n"):
                if para.strip():
                    doc.add_paragraph(para.strip())

        if overview.get("decision_point"):
            add_h2("Decision Point")
            dp = doc.add_paragraph()
            dr = dp.add_run(overview["decision_point"])
            dr.font.bold = True
            dr.font.color.rgb = RGBColor(20, 70, 140)

        if overview.get("key_individuals"):
            add_h2("Key Individuals")
            tbl = doc.add_table(rows=1, cols=3)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr = tbl.rows[0].cells
            hdr[0].paragraphs[0].add_run("Name").font.bold = True
            hdr[1].paragraphs[0].add_run("Title").font.bold = True
            hdr[2].paragraphs[0].add_run("Role in Case").font.bold = True
            for ind in overview["key_individuals"]:
                row = tbl.add_row().cells
                row[0].paragraphs[0].add_run(ind.get("name", ""))
                row[1].paragraphs[0].add_run(ind.get("title", ""))
                row[2].paragraphs[0].add_run(ind.get("role", ""))

        if overview.get("contextual_facts"):
            add_h2("Key Contextual Facts Students Need Front-of-Mind")
            for fact in overview["contextual_facts"]:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(fact)

        if overview.get("red_herrings"):
            add_h2("Red Herrings")
            for rh in overview["red_herrings"]:
                p = doc.add_paragraph(style='List Bullet')
                r = p.add_run(f"{rh.get('trap', '')}: ")
                r.font.bold = True
                p.add_run(rh.get('guidance', ''))

    # 2. LEARNING OBJECTIVES
    add_h1("2. Learning Objectives & Why Teach This Case")
    if learning_objs:
        add_h2("Learning Objectives")
        for obj in learning_objs:
            p = doc.add_paragraph()
            p.add_run(obj)

    if why_teach:
        add_h2("Why Teach This Case")
        if why_teach.get("overview"):
            doc.add_paragraph(why_teach["overview"])
        if why_teach.get("suitable_courses"):
            for c in why_teach["suitable_courses"]:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(c)
        if why_teach.get("setting_significance"):
            doc.add_paragraph(why_teach["setting_significance"])

    # 3. THEORY & FRAMEWORKS
    if theories:
        add_h1("3. Theory and Frameworks")
        for i, t in enumerate(theories, 1):
            add_h2(f"{i}. {t.get('framework_name', '')}")
            p = doc.add_paragraph()
            pr = p.add_run("Pedagogical purpose: ")
            pr.font.italic = True
            p.add_run(t.get("pedagogical_purpose", ""))

    # 4. CLASSROOM STRATEGY
    if strategy and strategy.get("activities"):
        add_h1("4. Classroom Strategy & In-Class Activities")
        for act in strategy["activities"]:
            add_h2(act.get("title", "Classroom Activity"))
            if act.get("setup"):
                p = doc.add_paragraph()
                p.add_run("Setup: ").font.bold = True
                p.add_run(act["setup"])
            if act.get("implementation"):
                p = doc.add_paragraph()
                p.add_run("Implementation: ").font.bold = True
                p.add_run(act["implementation"])
            if act.get("why_it_works"):
                p = doc.add_paragraph()
                p.add_run("Why it works: ").font.bold = True
                p.add_run(act["why_it_works"])

    # 5. DISCUSSION PLAN
    if disc_plan:
        add_h1("5. Discussion Plan")
        if disc_plan.get("assumed_format"):
            p = doc.add_paragraph()
            p.add_run(f"Assumed format: {disc_plan['assumed_format']}").font.bold = True

        pastures = disc_plan.get("pastures", [])
        for past in pastures:
            p_num = past.get("pasture_number", "")
            p_title = past.get("title", "")
            p_time = past.get("estimated_time", "")
            add_h2(f"Pasture {p_num}: {p_title} ({p_time})")

            if past.get("framing"):
                p = doc.add_paragraph()
                r = p.add_run("Framing for instructor: ")
                r.font.bold = True
                r.font.italic = True
                p.add_run(past["framing"])

            for q in past.get("questions", []):
                qp = doc.add_paragraph()
                qp.paragraph_format.space_before = Pt(6)
                qr = qp.add_run(f"Question {q.get('question_id', '')}\n")
                qr.font.bold = True
                qr.font.size = Pt(10.5)
                qr.font.color.rgb = RGBColor(15, 41, 74)
                q_text_run = qp.add_run(q.get("question_text", ""))
                q_text_run.font.bold = True

                if q.get("board_plan_matrix"):
                    bp_tbl = doc.add_table(rows=1, cols=3)
                    bp_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                    bp_hdr = bp_tbl.rows[0].cells
                    bp_hdr[0].paragraphs[0].add_run("Stakeholder").font.bold = True
                    bp_hdr[1].paragraphs[0].add_run("Nature of Trust Breach").font.bold = True
                    bp_hdr[2].paragraphs[0].add_run("Repair Mechanism").font.bold = True
                    for row_data in q["board_plan_matrix"]:
                        r_cells = bp_tbl.add_row().cells
                        r_cells[0].paragraphs[0].add_run(row_data.get("stakeholder", ""))
                        r_cells[1].paragraphs[0].add_run(row_data.get("nature_of_breach", ""))
                        r_cells[2].paragraphs[0].add_run(row_data.get("repair_mechanism", ""))

                exp = q.get("expected_responses", {})
                if exp:
                    ep = doc.add_paragraph()
                    ep.add_run("Expected responses:").font.bold = True
                    if exp.get("encourage"):
                        p_enc = doc.add_paragraph(style='List Bullet')
                        r_enc = p_enc.add_run("• Encourage: ")
                        r_enc.font.bold = True
                        r_enc.font.color.rgb = RGBColor(4, 120, 87)
                        p_enc.add_run(exp["encourage"])
                    if exp.get("redirect"):
                        p_red = doc.add_paragraph(style='List Bullet')
                        r_red = p_red.add_run("• Redirect: ")
                        r_red.font.bold = True
                        r_red.font.color.rgb = RGBColor(180, 50, 20)
                        p_red.add_run(exp["redirect"])
                    if exp.get("if_struggling"):
                        p_st = doc.add_paragraph()
                        p_st.add_run("If students struggle: ").font.bold = True
                        p_st.add_run(exp["if_struggling"])

                if q.get("classroom_management_note"):
                    p_cm = doc.add_paragraph()
                    p_cm.add_run("Classroom management note: ").font.bold = True
                    p_cm.add_run(q["classroom_management_note"])

                if q.get("analysis"):
                    p_an = doc.add_paragraph()
                    p_an.add_run("Analysis: ").font.bold = True
                    p_an.add_run(q["analysis"])

                if q.get("framework_connection"):
                    p_fc = doc.add_paragraph()
                    p_fc.add_run("Framework connection: ").font.bold = True
                    p_fc.add_run(q["framework_connection"])

                if q.get("estimated_time"):
                    p_et = doc.add_paragraph()
                    p_et.add_run(f"Estimated time: {q['estimated_time']}").font.italic = True

        if disc_plan.get("timing_summary"):
            add_h2("Timing Summary")
            t_tbl = doc.add_table(rows=1, cols=3)
            t_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            th = t_tbl.rows[0].cells
            th[0].paragraphs[0].add_run("Pasture").font.bold = True
            th[1].paragraphs[0].add_run("Focus").font.bold = True
            th[2].paragraphs[0].add_run("Time").font.bold = True
            for tr_item in disc_plan["timing_summary"]:
                row_c = t_tbl.add_row().cells
                row_c[0].paragraphs[0].add_run(tr_item.get("pasture", ""))
                row_c[1].paragraphs[0].add_run(tr_item.get("focus", ""))
                row_c[2].paragraphs[0].add_run(tr_item.get("time", ""))

    # 6. TAKEAWAYS & APPLICATION
    if takeaways:
        add_h1("6. Takeaways & Application")
        if takeaways.get("strategy_and_consulting"):
            add_h2("For Students Entering Strategy, Consulting, or Corporate Development Roles")
            for t in takeaways["strategy_and_consulting"]:
                doc.add_paragraph(t, style='List Bullet')
        if takeaways.get("general_management_and_leadership"):
            add_h2("For Students Entering General Management or Leadership Roles")
            for t in takeaways["general_management_and_leadership"]:
                doc.add_paragraph(t, style='List Bullet')
        if takeaways.get("ethics_and_governance"):
            add_h2("For Students Interested in Ethics and Corporate Governance")
            for t in takeaways["ethics_and_governance"]:
                doc.add_paragraph(t, style='List Bullet')

    # 7. ALTERNATIVE APPROACHES
    if alternatives:
        add_h1("7. Alternative Approaches")
        for alt in alternatives:
            add_h2(alt.get("approach_name", "Alternative Approach"))
            doc.add_paragraph(alt.get("description", ""))

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)

    filename = f"AI_Case_Note_{safe_title}.docx"
    return StreamingResponse(
        bio,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )


def _generate_case_note_html_response(case_note: dict, case_title: str) -> HTMLResponse:
    import html
    def esc(text) -> str:
        return html.escape(str(text or ""))

    meta = case_note.get("meta", {})
    overview = case_note.get("overview", {})
    learning_objs = case_note.get("learning_objectives", [])
    why_teach = case_note.get("why_teach_this_case", {})
    theories = case_note.get("theory_and_frameworks", [])
    strategy = case_note.get("classroom_strategy", {})
    disc_plan = case_note.get("discussion_plan", {})
    takeaways = case_note.get("takeaways_and_application", {})
    alternatives = case_note.get("alternative_approaches", [])

    auth_name = esc(meta.get("authorized_for", "Faculty / Student, Lexicon MILE"))
    gen_time = esc(meta.get("generated_timestamp", ""))
    case_code = esc(meta.get("case_code", "HBP-CASE"))
    authors = esc(meta.get("authors", ""))
    setting = esc(meta.get("setting", ""))
    duration = esc(meta.get("assumed_duration", "80-minute in-person class session"))

    html_parts = [f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>AI Case Note — {esc(meta.get('case_title', case_title))}</title>
<style>
  @page {{
    size: letter;
    margin: 1.8cm 1.8cm 1.8cm 1.8cm;
  }}
  body {{
    font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", Arial, sans-serif;
    color: #1e293b;
    line-height: 1.6;
    background: #f8fafc;
    margin: 0;
    padding: 24px;
    font-size: 10pt;
  }}
  .document-container {{
    max-width: 860px;
    margin: 0 auto;
    background: #ffffff;
    padding: 48px 56px;
    border: 1px solid #e2e8f0;
    box-shadow: 0 4px 20px rgba(0,0,0,0.06);
    border-radius: 4px;
  }}
  .hbp-top-banner {{
    border-bottom: 2px solid #0f172a;
    padding-bottom: 12px;
    margin-bottom: 24px;
  }}
  .hbp-dnc {{
    font-size: 9pt;
    font-weight: 800;
    text-transform: uppercase;
    letter-spacing: 0.05em;
    color: #b91c1c;
    margin-bottom: 2px;
  }}
  .hbp-created-for {{
    font-size: 8.5pt;
    color: #64748b;
    font-style: italic;
  }}
  .hbp-main-title {{
    font-family: "Georgia", serif;
    font-size: 22pt;
    font-weight: 900;
    color: #0f294a;
    margin: 16px 0 8px 0;
    letter-spacing: -0.01em;
    line-height: 1.2;
  }}
  .hbp-supporting {{
    font-size: 11pt;
    font-weight: 700;
    color: #334155;
    margin-bottom: 14px;
  }}
  .hbp-supporting-sub {{
    font-size: 10pt;
    font-weight: normal;
    color: #475569;
  }}
  .hbp-auth-box {{
    background: #f1f5f9;
    border-left: 3px solid #0f294a;
    padding: 10px 14px;
    font-size: 8.5pt;
    color: #475569;
    margin: 14px 0 20px 0;
    line-height: 1.5;
  }}
  .hbp-meta-bar {{
    font-size: 9pt;
    font-weight: 600;
    color: #0f294a;
    padding: 6px 0;
    border-top: 1px solid #cbd5e1;
    border-bottom: 1px solid #cbd5e1;
    margin-bottom: 24px;
  }}
  h1.sec-h1 {{
    font-family: "Georgia", serif;
    font-size: 14pt;
    font-weight: 800;
    color: #0f294a;
    border-bottom: 1.5px solid #0f294a;
    padding-bottom: 4px;
    margin-top: 28px;
    margin-bottom: 14px;
    page-break-after: avoid;
  }}
  h2.sec-h2 {{
    font-size: 11pt;
    font-weight: 700;
    color: #1e3a8a;
    margin-top: 18px;
    margin-bottom: 6px;
    page-break-after: avoid;
  }}
  p {{
    margin: 0 0 10px 0;
    text-align: justify;
  }}
  .decision-box {{
    background: #f8fafc;
    border: 1px solid #cbd5e1;
    border-left: 4px solid #2563eb;
    padding: 12px 16px;
    margin: 12px 0;
    font-weight: 600;
    color: #0f172a;
    border-radius: 2px;
  }}
  table {{
    width: 100%;
    border-collapse: collapse;
    margin: 14px 0;
    font-size: 9pt;
  }}
  th {{
    background: #0f294a;
    color: #ffffff;
    font-weight: 700;
    text-align: left;
    padding: 8px 10px;
    border: 1px solid #0f294a;
  }}
  td {{
    padding: 7px 10px;
    border: 1px solid #e2e8f0;
    vertical-align: top;
  }}
  tr:nth-child(even) td {{
    background: #f8fafc;
  }}
  ul, ol {{
    margin: 6px 0 14px 0;
    padding-left: 22px;
  }}
  li {{
    margin-bottom: 5px;
  }}
  .q-block {{
    background: #ffffff;
    border: 1px solid #e2e8f0;
    border-left: 3px solid #0f294a;
    padding: 14px 16px;
    margin: 14px 0;
    page-break-inside: avoid;
  }}
  .q-num {{
    font-weight: 800;
    color: #0f294a;
    font-size: 10.5pt;
    margin-bottom: 4px;
  }}
  .q-text {{
    font-weight: 700;
    font-style: italic;
    color: #0f172a;
    margin-bottom: 10px;
    font-size: 10pt;
  }}
  .enc-tag {{
    color: #047857;
    font-weight: 800;
  }}
  .red-tag {{
    color: #b91c1c;
    font-weight: 800;
  }}
  .struggle-box {{
    background: #fffbeb;
    border: 1px solid #fef3c7;
    padding: 6px 10px;
    font-size: 8.5pt;
    color: #92400e;
    margin: 8px 0;
    border-radius: 2px;
  }}
  .mgmt-box {{
    background: #fdf2f8;
    border: 1px solid #fce7f3;
    padding: 6px 10px;
    font-size: 8.5pt;
    color: #9d174d;
    margin: 8px 0;
    border-radius: 2px;
  }}
  .analysis-text {{
    font-size: 9pt;
    color: #334155;
    margin-top: 8px;
    padding-top: 6px;
    border-top: 1px dashed #cbd5e1;
  }}
  .activity-card {{
    background: #f8fafc;
    border: 1px solid #e2e8f0;
    padding: 12px 16px;
    margin: 10px 0;
    border-radius: 3px;
    page-break-inside: avoid;
  }}
  .footer-bar {{
    margin-top: 36px;
    padding-top: 12px;
    border-top: 1px solid #cbd5e1;
    font-size: 8pt;
    color: #64748b;
    display: flex;
    justify-content: space-between;
  }}
  @media print {{
    body {{
      background: #ffffff;
      padding: 0;
      font-size: 9.5pt;
    }}
    .document-container {{
      border: none;
      box-shadow: none;
      padding: 0;
      max-width: 100%;
    }}
    .no-print {{
      display: none !important;
    }}
  }}
</style>
</head>
<body>

<div class="no-print" style="max-width:860px;margin:0 auto 16px auto;display:flex;justify-content:space-between;align-items:center;padding:10px 16px;background:#0f294a;color:#fff;border-radius:6px;">
  <div><strong>AI Case Note View</strong> — Authorized Institutional Copy</div>
  <button onclick="window.print()" style="background:#2563eb;color:#fff;border:none;padding:6px 14px;border-radius:4px;font-weight:bold;cursor:pointer;">🖨️ Print / Save as PDF</button>
</div>

<div class="document-container">
  <div class="hbp-top-banner">
    <div class="hbp-dnc">Do Not Copy or Post</div>
    <div class="hbp-created-for">Case Note created for {auth_name} on {gen_time}</div>
    <div class="hbp-main-title">AI Case Note — {esc(meta.get('case_title', case_title))}</div>
    <div class="hbp-supporting">SUPPORTING:<br><span class="hbp-supporting-sub">{esc(meta.get('case_title', case_title))}<br>By {authors}</span></div>
    <div class="hbp-auth-box">
      This Case Note is authorized for <strong>{auth_name}</strong>.<br>
      <em>This case note was created by generative AI. Although we strive for accuracy, its responses may be inaccurate or incomplete. Please review all information provided carefully.</em>
    </div>
    <div class="hbp-meta-bar">
      Case: {case_code} &nbsp;|&nbsp; Authors: {authors} &nbsp;|&nbsp; Setting: {setting}
    </div>
  </div>
"""]

    if overview:
        html_parts.append('<h1 class="sec-h1">1. Case Overview</h1>')
        if overview.get("story_synopsis"):
            html_parts.append('<h2 class="sec-h2">Story Synopsis</h2>')
            for p in overview["story_synopsis"].split("\n\n"):
                if p.strip():
                    html_parts.append(f'<p>{esc(p.strip())}</p>')

        if overview.get("decision_point"):
            html_parts.append('<h2 class="sec-h2">Decision Point</h2>')
            html_parts.append(f'<div class="decision-box">{esc(overview["decision_point"])}</div>')

        if overview.get("key_individuals"):
            html_parts.append('<h2 class="sec-h2">Key Individuals</h2>')
            html_parts.append('<table><tr><th style="width:25%;">Name</th><th style="width:30%;">Title</th><th>Role in Case</th></tr>')
            for ind in overview["key_individuals"]:
                html_parts.append(f'<tr><td><strong>{esc(ind.get("name",""))}</strong></td><td>{esc(ind.get("title",""))}</td><td>{esc(ind.get("role",""))}</td></tr>')
            html_parts.append('</table>')

        if overview.get("contextual_facts"):
            html_parts.append('<h2 class="sec-h2">Key Contextual Facts Students Need Front-of-Mind</h2><ul>')
            for fact in overview["contextual_facts"]:
                html_parts.append(f'<li>{esc(fact)}</li>')
            html_parts.append('</ul>')

        if overview.get("red_herrings"):
            html_parts.append('<h2 class="sec-h2">Red Herrings</h2><ul>')
            for rh in overview["red_herrings"]:
                html_parts.append(f'<li><strong>{esc(rh.get("trap",""))}:</strong> {esc(rh.get("guidance",""))}</li>')
            html_parts.append('</ul>')

    html_parts.append('<h1 class="sec-h1">2. Learning Objectives & Why Teach This Case</h1>')
    if learning_objs:
        html_parts.append('<h2 class="sec-h2">Learning Objectives</h2><p><em>By the end of class, students should be able to:</em></p><ol>')
        for obj in learning_objs:
            html_parts.append(f'<li>{esc(obj)}</li>')
        html_parts.append('</ol>')

    if why_teach:
        html_parts.append('<h2 class="sec-h2">Why Teach This Case</h2>')
        if why_teach.get("overview"):
            html_parts.append(f'<p>{esc(why_teach["overview"])}</p>')
        if why_teach.get("suitable_courses"):
            html_parts.append('<ul>')
            for c in why_teach["suitable_courses"]:
                html_parts.append(f'<li>{esc(c)}</li>')
            html_parts.append('</ul>')
        if why_teach.get("setting_significance"):
            html_parts.append(f'<p>{esc(why_teach["setting_significance"])}</p>')

    if theories:
        html_parts.append('<h1 class="sec-h1">3. Theory and Frameworks</h1><ol>')
        for t in theories:
            html_parts.append(f'<li><strong>{esc(t.get("framework_name",""))}</strong><br><span style="font-size:9pt;color:#334155;"><em>Pedagogical purpose:</em> {esc(t.get("pedagogical_purpose",""))}</span></li>')
        html_parts.append('</ol>')

    if strategy and strategy.get("activities"):
        html_parts.append('<h1 class="sec-h1">4. Classroom Strategy</h1>')
        for act in strategy["activities"]:
            html_parts.append(f"""<div class="activity-card">
              <h3 style="margin:0 0 6px 0;font-size:10.5pt;color:#0f294a;">{esc(act.get("title",""))}</h3>
              <p><strong>Setup:</strong> {esc(act.get("setup",""))}</p>
              <p><strong>Implementation:</strong> {esc(act.get("implementation",""))}</p>
              <p style="margin:0;"><strong>Why it works:</strong> {esc(act.get("why_it_works",""))}</p>
            </div>""")

    if disc_plan:
        html_parts.append('<h1 class="sec-h1">5. Discussion Plan</h1>')
        if disc_plan.get("assumed_format"):
            html_parts.append(f'<p><strong>Assumed format:</strong> {esc(disc_plan["assumed_format"])}</p>')

        pastures = disc_plan.get("pastures", [])
        for past in pastures:
            p_num = esc(past.get("pasture_number", ""))
            p_title = esc(past.get("title", ""))
            p_time = esc(past.get("estimated_time", ""))
            html_parts.append(f'<h2 class="sec-h2" style="font-size:12pt;border-top:1px solid #e2e8f0;padding-top:10px;">Pasture {p_num}: {p_title} ({p_time})</h2>')
            if past.get("framing"):
                html_parts.append(f'<p><em><strong>Framing for instructor:</strong> {esc(past["framing"])}</em></p>')

            for q in past.get("questions", []):
                q_id = esc(q.get("question_id", ""))
                q_text = esc(q.get("question_text", ""))
                html_parts.append(f'<div class="q-block"><div class="q-num">Question {q_id}</div><div class="q-text">{q_text}</div>')

                if q.get("board_plan_matrix"):
                    html_parts.append('<p style="font-size:8.5pt;font-weight:bold;margin:6px 0 2px 0;">Board plan: Draw a simple matrix on the board as students respond:</p>')
                    html_parts.append('<table><tr><th>Stakeholder</th><th>Nature of Trust Breach</th><th>Repair Mechanism</th></tr>')
                    for row_data in q["board_plan_matrix"]:
                        html_parts.append(f'<tr><td><strong>{esc(row_data.get("stakeholder",""))}</strong></td><td>{esc(row_data.get("nature_of_breach",""))}</td><td>{esc(row_data.get("repair_mechanism",""))}</td></tr>')
                    html_parts.append('</table>')

                exp = q.get("expected_responses", {})
                if exp:
                    html_parts.append('<div style="font-size:9pt;font-weight:700;margin-top:8px;">Expected responses:</div><ul>')
                    if exp.get("encourage"):
                        html_parts.append(f'<li><span class="enc-tag">• Encourage:</span> {esc(exp["encourage"])}</li>')
                    if exp.get("redirect"):
                        html_parts.append(f'<li><span class="red-tag">• Redirect:</span> {esc(exp["redirect"])}</li>')
                    html_parts.append('</ul>')
                    if exp.get("if_struggling"):
                        html_parts.append(f'<div class="struggle-box"><strong>If students struggle:</strong> {esc(exp["if_struggling"])}</div>')

                if q.get("classroom_management_note"):
                    html_parts.append(f'<div class="mgmt-box"><strong>Classroom management note:</strong> {esc(q["classroom_management_note"])}</div>')

                if q.get("analysis"):
                    html_parts.append(f'<div class="analysis-text"><strong>Analysis:</strong> {esc(q["analysis"])}</div>')

                if q.get("framework_connection"):
                    html_parts.append(f'<div style="font-size:8.5pt;color:#1e3a8a;margin-top:4px;"><strong>Framework connection:</strong> {esc(q["framework_connection"])}</div>')

                if q.get("estimated_time"):
                    html_parts.append(f'<div style="font-size:8pt;color:#64748b;font-style:italic;margin-top:4px;">Estimated time: {esc(q["estimated_time"])}</div>')

                html_parts.append('</div>')

        if disc_plan.get("timing_summary"):
            html_parts.append('<h2 class="sec-h2">Timing Summary</h2><table><tr><th>Pasture</th><th>Focus</th><th>Time</th></tr>')
            for tr_item in disc_plan["timing_summary"]:
                html_parts.append(f'<tr><td><strong>{esc(tr_item.get("pasture",""))}</strong></td><td>{esc(tr_item.get("focus",""))}</td><td>{esc(tr_item.get("time",""))}</td></tr>')
            html_parts.append('</table>')

    if takeaways:
        html_parts.append('<h1 class="sec-h1">6. Takeaways & Application</h1>')
        if takeaways.get("strategy_and_consulting"):
            html_parts.append('<h2 class="sec-h2">For Students Entering Strategy, Consulting, or Corporate Development Roles</h2><ul>')
            for t in takeaways["strategy_and_consulting"]:
                html_parts.append(f'<li>{esc(t)}</li>')
            html_parts.append('</ul>')
        if takeaways.get("general_management_and_leadership"):
            html_parts.append('<h2 class="sec-h2">For Students Entering General Management or Leadership Roles</h2><ul>')
            for t in takeaways["general_management_and_leadership"]:
                html_parts.append(f'<li>{esc(t)}</li>')
            html_parts.append('</ul>')
        if takeaways.get("ethics_and_governance"):
            html_parts.append('<h2 class="sec-h2">For Students Interested in Ethics and Corporate Governance</h2><ul>')
            for t in takeaways["ethics_and_governance"]:
                html_parts.append(f'<li>{esc(t)}</li>')
            html_parts.append('</ul>')

    if alternatives:
        html_parts.append('<h1 class="sec-h1">7. Alternative Approaches</h1>')
        for alt in alternatives:
            html_parts.append(f'<h2 class="sec-h2">{esc(alt.get("approach_name",""))}</h2><p>{esc(alt.get("description",""))}</p>')

    html_parts.append(f"""
    <div class="footer-bar">
      <div>AI Case Note — {esc(meta.get('case_title', case_title))}</div>
      <div>Authorized for {auth_name}</div>
    </div>
  </div>
</body>
</html>""")

    return HTMLResponse(content="".join(html_parts))
