import io
import re
import json
import uuid
import logging
import asyncio
from datetime import datetime
from typing import Dict, Any, List, Optional, Tuple
from uuid import UUID

import httpx
from pypdf import PdfReader
from docx import Document
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import settings
from app.models.activity import SubjectActivity
from app.models.academic import Subject

logger = logging.getLogger(__name__)

NEON_URI = "postgresql://neondb_owner:npg_qKgko53pGtyx@ep-proud-poetry-aze89kl6-pooler.c-3.ap-southeast-1.aws.neon.tech/neondb"


def extract_text_from_doc_bytes(data: bytes, filename: str) -> Tuple[str, List[str]]:
    """
    Extracts text from PDF or DOCX file bytes.
    Returns (full_text, pages_or_sections).
    """
    ext = (filename or "").lower().split(".")[-1]
    pages = []

    if ext == "pdf":
        reader = PdfReader(io.BytesIO(data))
        for i, page in enumerate(reader.pages):
            txt = page.extract_text() or ""
            pages.append(f"--- Page {i+1} ---\n{txt}")
        full_text = "\n\n".join(pages)
        return full_text, pages

    elif ext in ["docx", "doc"]:
        doc = Document(io.BytesIO(data))
        paragraphs = []
        for p in doc.paragraphs:
            if p.text.strip():
                paragraphs.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_str = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_str:
                    paragraphs.append(f"| {row_str} |")
        full_text = "\n".join(paragraphs)
        return full_text, [full_text]

    else:
        txt = data.decode("utf-8", errors="ignore")
        return txt, [txt]


def detect_subject_from_text(text: str) -> Tuple[Optional[str], Optional[str]]:
    """
    Detects subject name and subject code from document headers.
    e.g. 'Statistics for Managers (1BAC01) | PGDM Batch 2026–28'
    """
    m_code = re.search(r'([A-Za-z\s&,]+?)\s*\(([A-Z0-9]{5,8})\)', text)
    if m_code:
        return m_code.group(1).strip(), m_code.group(2).strip()

    m_header = re.search(r'HyperBuild\s*(?:AI\s*Lab)?\s*\|\s*([^\n|]+)', text, re.IGNORECASE)
    if m_header:
        candidate = m_header.group(1).strip()
        if len(candidate) > 3 and not candidate.startswith("Page"):
            return candidate, None

    return None, None


def split_into_activity_chunks(text: str) -> List[str]:
    """
    Splits multi-activity document text into individual activity sections.
    Handles single-activity documents and master documents containing up to 16 activities.
    """
    boundary_pattern = re.compile(
        r'(?=(?:--- Page \d+ ---\s*)?(?:Activity\s+(\d+)\s*[:—\-]|ACTIVITY\s+(\d+)|CAPSTONE\s+PROJECT))',
        re.IGNORECASE
    )

    matches = list(boundary_pattern.finditer(text))
    if not matches:
        return [text]

    deduped_starts = []
    seen_activities = set()

    for m in matches:
        pos = m.start()
        snippet = text[pos:pos+250]
        act_match = re.search(r'(?:Activity\s+(\d+)|ACTIVITY\s+(\d+)|CAPSTONE)', snippet, re.IGNORECASE)
        act_num = None
        if act_match:
            if "CAPSTONE" in snippet.upper():
                act_num = 16
            else:
                act_num = int(act_match.group(1) or act_match.group(2))

        if act_num is not None:
            if act_num not in seen_activities:
                seen_activities.add(act_num)
                deduped_starts.append(pos)
        else:
            if not deduped_starts or pos - deduped_starts[-1] > 500:
                deduped_starts.append(pos)

    if not deduped_starts:
        return [text]

    chunks = []
    for i, start_pos in enumerate(deduped_starts):
        end_pos = deduped_starts[i+1] if i + 1 < len(deduped_starts) else len(text)
        chunk = text[start_pos:end_pos].strip()
        if len(chunk) > 200:
            chunks.append(chunk)

    return chunks if chunks else [text]


EXTRACTION_PROMPT = """You are an expert curriculum data extraction system for a premier business school LMS.
Parse the following HyperBuild course activity document section into a valid, well-structured JSON object.

Document section text:
\"\"\"
{chunk_text}
\"\"\"

Return a single JSON object with this exact schema:
{{
  "activity_no": 5,
  "title": "Exact full title of the activity (e.g. Bayes' Theorem & Fraud/Credit Risk Lab — Why a 95%-Accurate Model Can Still Be Wrong Most of the Time)",
  "unit_mapping": "Unit 2 — Probability, Risk, and Uncertainty in Business Decisions",
  "estimated_time": "3–4 hours",
  "mode": "Individual",
  "file_naming": "[RollNo]_SFM_Act05",
  "why_this_activity": "• Bullet point 1...\\n• Bullet point 2...",
  "instructions": "1. Step 1...\\n2. Step 2...",
  "ai_tools": [
    {{
      "tool": "Microsoft Excel",
      "category": "Bayesian Modelling",
      "purpose": "Law of Total Probability + Bayes' Theorem computation for two scenarios",
      "access": "Office 365 / Installed"
    }}
  ],
  "learning_outcomes": "• Outcome 1...\\n• Outcome 2...",
  "submission_requirements": "• (1) Excel workbook...\\n• (2) 350-word Memo...",
  "rubric": [
    {{
      "criterion": "Criterion Name",
      "weightage": 40,
      "distinction": "Detailed Distinction 85-100% criteria descriptor",
      "merit": "Detailed Merit 70-84% criteria descriptor",
      "pass_grade": "Detailed Pass 50-69% criteria descriptor",
      "needs_work": "Detailed Needs Work <50% criteria descriptor"
    }}
  ]
}}

CRITICAL RULES:
1. Return ONLY pure valid JSON without markdown code block backticks (like ```json), commentary, or notes.
2. Ensure rubric weights are integers that sum up to 100.
3. Preserve all mathematical equations, scenario numbers, and student instructions exactly as written.
4. Extract all AI tools mentioned in the AI TOOLS & PLATFORMS table.
"""


async def parse_activity_chunk_with_ai(chunk_text: str) -> Optional[Dict[str, Any]]:
    """
    Parses an individual activity text chunk into structured JSON using Groq LPU with OpenRouter fallback.
    """
    prompt = EXTRACTION_PROMPT.format(chunk_text=chunk_text[:18000])

    # 1. Try Groq with openai/gpt-oss-120b
    groq_key = settings.GROQ_API_KEY
    if groq_key:
        try:
            url = "https://api.groq.com/openai/v1/chat/completions"
            headers = {"Authorization": f"Bearer {groq_key}", "Content-Type": "application/json"}
            payload = {
                "model": "openai/gpt-oss-120b",
                "messages": [{"role": "user", "content": prompt}],
                "response_format": {"type": "json_object"},
                "temperature": 0.1,
                "max_tokens": 4000
            }
            async with httpx.AsyncClient(timeout=45.0) as client:
                res = await client.post(url, headers=headers, json=payload)
                if res.status_code == 200:
                    data = res.json()["choices"][0]["message"]["content"]
                    return json.loads(data)
                else:
                    logger.warning(f"Groq activity extraction returned {res.status_code}: {res.text[:150]}")
        except Exception as e:
            logger.warning(f"Groq activity extraction failed: {e}")

    # 2. Try OpenRouter fallback with llama-3.3-70b-instruct
    or_key = settings.OPENROUTER_API_KEY
    if or_key:
        try:
            or_url = "https://openrouter.ai/api/v1/chat/completions"
            or_headers = {"Authorization": f"Bearer {or_key}", "Content-Type": "application/json"}
            or_payload = {
                "model": "meta-llama/llama-3.3-70b-instruct",
                "messages": [{"role": "user", "content": prompt}],
                "response_format": {"type": "json_object"},
                "temperature": 0.1
            }
            async with httpx.AsyncClient(timeout=45.0) as client:
                res = await client.post(or_url, headers=or_headers, json=or_payload)
                if res.status_code == 200:
                    raw = res.json()["choices"][0]["message"]["content"]
                    if raw.startswith("```"):
                        raw = raw.split("\n", 1)[1].rsplit("```", 1)[0].strip()
                    return json.loads(raw)
        except Exception as e:
            logger.warning(f"OpenRouter activity extraction failed: {e}")

    # 3. Deterministic regex fallback
    return parse_activity_chunk_deterministic(chunk_text)


def parse_activity_chunk_deterministic(text: str) -> Optional[Dict[str, Any]]:
    """
    Deterministic regex fallback parser for HyperBuild activity structure.
    """
    act_no_match = re.search(r'(?:Activity\s+(\d+)|ACTIVITY\s+(\d+))', text, re.IGNORECASE)
    act_no = int(act_no_match.group(1) or act_no_match.group(2)) if act_no_match else None
    if not act_no and "CAPSTONE" in text.upper():
        act_no = 16

    if not act_no:
        return None

    # Title
    title = ""
    title_m = re.search(r'ACTIVITY\s*\d+\s*\n+([^\n]+(?:\n+[^\n]+)?)\s*\n+Unit\s+\d+', text, re.IGNORECASE)
    if title_m:
        title = " ".join(l.strip() for l in title_m.group(1).split("\n") if l.strip())
    else:
        alt_m = re.search(r'(?:Activity\s+\d+\s*[:—\-]\s*)([^\n]+)', text, re.IGNORECASE)
        title = alt_m.group(1).strip() if alt_m else f"Activity {act_no}"

    unit_mapping = ""
    unit_m = re.search(r'(Unit\s+\d+\s*[—\-–][^\n]+)', text, re.IGNORECASE)
    if unit_m:
        unit_mapping = unit_m.group(1).strip()

    est_time = "3 hours"
    time_m = re.search(r'Estimated Time:\s*([^\n|]+?)(?:\s+Mode:|$)', text, re.IGNORECASE)
    if time_m:
        est_time = time_m.group(1).strip()

    mode = "Individual"
    mode_m = re.search(r'Mode:\s*([^\n|]+?)(?:\s+File Naming:|$)', text, re.IGNORECASE)
    if mode_m:
        mode = mode_m.group(1).strip()

    file_naming = f"[RollNo]_Act{act_no:02d}"
    file_m = re.search(r'File Naming:\s*([^\n|]+)', text, re.IGNORECASE)
    if file_m:
        file_naming = file_m.group(1).strip()

    def extract_sec(start_kw, end_kws):
        pattern = rf'{start_kw}\s*\n+(.*?)(?=(?:' + '|'.join(end_kws) + r')|\Z)'
        m = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
        if not m:
            return ""
        c = re.sub(r'--- Page \d+ ---\s*HyperBuild[^\n]*\n[^\n]*\n', '', m.group(1))
        c = re.sub(r'HyperBuild[^\n]*\n[^\n]*\n', '', c).strip()
        return c

    why = extract_sec(r'WHY THIS ACTIVITY', [r'STEP-BY-STEP INSTRUCTIONS', r'CAPSTONE STRUCTURE', r'AI TOOLS'])
    instructions = extract_sec(r'(?:STEP-BY-STEP INSTRUCTIONS|CAPSTONE STRUCTURE)', [r'AI TOOLS', r'LEARNING OUTCOMES'])
    outcomes = extract_sec(r'LEARNING OUTCOMES', [r'ASSESSMENT TASK', r'GRADING RUBRIC'])
    reqs = extract_sec(r'(?:ASSESSMENT TASK & SUBMISSION REQUIREMENTS|ASSESSMENT TASK)', [r'GRADING RUBRIC'])

    rubric = [
        {
            "criterion": "Technical & Analytical Execution",
            "weightage": 40,
            "distinction": "Exceptional formulation, rigorous computational model, and clear business justification.",
            "merit": "Accurate computation and sound application with minor presentation gaps.",
            "pass_grade": "Basic requirements attempted with partial computational correctness.",
            "needs_work": "Substantial errors or incomplete deliverable."
        },
        {
            "criterion": "AI Cross-Verification & Critical Synthesis",
            "weightage": 30,
            "distinction": "Thorough comparison between AI outputs and student calculations with insightful resolution.",
            "merit": "AI comparison performed with adequate student synthesis.",
            "pass_grade": "AI interaction documented without deep comparative analysis.",
            "needs_work": "AI verification missing or uncritical acceptance of generated outputs."
        },
        {
            "criterion": "Executive Memo & Strategic Interpretation",
            "weightage": 30,
            "distinction": "Persuasive, data-backed strategic memo delivering actionable operational decisions.",
            "merit": "Clear managerial summary communicating core findings effectively.",
            "pass_grade": "General discussion lacking quantitative precision or actionable depth.",
            "needs_work": "Memo absent or below required word count."
        }
    ]

    return {
        "activity_no": act_no,
        "title": title,
        "unit_mapping": unit_mapping,
        "estimated_time": est_time,
        "mode": mode,
        "file_naming": file_naming,
        "why_this_activity": why,
        "instructions": instructions,
        "ai_tools": [
            {
                "tool": "Microsoft Excel",
                "category": "Analysis & Modelling",
                "purpose": "Formulas and scenario models",
                "access": "Office 365"
            },
            {
                "tool": "ChatGPT / Claude",
                "category": "LLM — Cross-Verification",
                "purpose": "Independent calculation verification and synthesis",
                "access": "Web Portal"
            }
        ],
        "learning_outcomes": outcomes,
        "submission_requirements": reqs,
        "rubric": rubric
    }


def normalize_activity_payload(act_data: Dict[str, Any]) -> Dict[str, Any]:
    """
    Validates, cleans, and normalizes activity payload before database insertion.
    Ensures rubric weights sum to 100.
    """
    act_no = int(act_data.get("activity_no", 1))
    title = str(act_data.get("title", f"Activity {act_no}")).strip()
    title = re.sub(r'^(?:Activity\s*\d+\s*[:—\-]*|ACTIVITY\s*\d+\s*[:—\-]*)\s*', '', title, flags=re.IGNORECASE).strip()

    rubric = act_data.get("rubric") or []
    if isinstance(rubric, str):
        try:
            rubric = json.loads(rubric)
        except Exception:
            rubric = []

    if rubric:
        total_wt = sum(int(r.get("weightage") or 0) for r in rubric)
        if total_wt != 100 and total_wt > 0:
            scale = 100.0 / total_wt
            running = 0
            for i, r in enumerate(rubric):
                if i == len(rubric) - 1:
                    r["weightage"] = 100 - running
                else:
                    new_w = round(int(r.get("weightage", 0)) * scale)
                    r["weightage"] = new_w
                    running += new_w

    return {
        "activity_no": act_no,
        "title": title,
        "unit_mapping": act_data.get("unit_mapping") or "",
        "estimated_time": act_data.get("estimated_time") or "3–4 hours",
        "mode": act_data.get("mode") or "Individual",
        "file_naming": act_data.get("file_naming") or f"[RollNo]_Act{act_no:02d}",
        "why_this_activity": act_data.get("why_this_activity") or "",
        "instructions": act_data.get("instructions") or "",
        "ai_tools": act_data.get("ai_tools") or [],
        "learning_outcomes": act_data.get("learning_outcomes") or "",
        "submission_requirements": act_data.get("submission_requirements") or "",
        "rubric": rubric,
        "is_released": True,
    }


async def import_activities_from_document(
    db: AsyncSession,
    subject_id: Optional[UUID],
    file_bytes: bytes,
    filename: str,
    overwrite_existing: bool = False
) -> Dict[str, Any]:
    """
    Main entry point: Ingests activities from a PDF or DOCX file into the database.
    - Preserves existing activities and student submissions by default.
    - Synchronizes new activities to both Aiven and Neon databases.
    """
    # 1. Extract text
    full_text, pages = extract_text_from_doc_bytes(file_bytes, filename)
    if not full_text or len(full_text.strip()) < 50:
        raise ValueError("Could not extract readable text from document. File may be empty or scanned.")

    # 2. Resolve Subject ID if not provided
    target_subject = None
    if subject_id:
        target_subject = await db.get(Subject, subject_id)
        if not target_subject:
            raise ValueError(f"Subject with ID '{subject_id}' not found.")
    else:
        # Fetch all subjects to match against document content
        res = await db.execute(select(Subject))
        all_subjects = list(res.scalars().all())
        full_text_lower = full_text.lower()

        # Match by subject name (longer names first)
        for s in sorted(all_subjects, key=lambda x: len(x.name), reverse=True):
            if s.name.lower() in full_text_lower:
                target_subject = s
                break

        # Fallback: match by code
        if not target_subject:
            for s in all_subjects:
                if s.code and s.code.lower() in full_text_lower:
                    target_subject = s
                    break

    if not target_subject:
        det_name, det_code = detect_subject_from_text(full_text)
        raise ValueError(
            f"Could not automatically determine subject from document (detected: '{det_name}' / '{det_code}'). "
            "Please select the target subject explicitly."
        )

    subj_id = target_subject.id

    # 3. Check existing activities in database for this subject
    existing_stmt = select(SubjectActivity).where(SubjectActivity.subject_id == subj_id)
    existing_res = await db.execute(existing_stmt)
    existing_acts = {a.activity_no: a for a in existing_res.scalars().all()}

    # 4. Split document into activity chunks
    chunks = split_into_activity_chunks(full_text)
    logger.info(f"Identified {len(chunks)} activity section(s) in {filename} for {target_subject.name}.")

    created_list = []
    skipped_list = []
    updated_list = []

    # 5. Process each activity section
    for idx, chunk in enumerate(chunks):
        parsed = await parse_activity_chunk_with_ai(chunk)
        if not parsed:
            continue

        clean = normalize_activity_payload(parsed)
        act_no = clean["activity_no"]

        if act_no in existing_acts:
            if not overwrite_existing:
                existing = existing_acts[act_no]
                skipped_list.append({
                    "activity_no": act_no,
                    "title": existing.title,
                    "status": "skipped",
                    "reason": f"Activity #{act_no} already exists in LMS. Preserved existing manual configuration and student submissions."
                })
                continue
            else:
                existing = existing_acts[act_no]
                for k, v in clean.items():
                    setattr(existing, k, v)
                existing.updated_at = datetime.utcnow()
                updated_list.append({
                    "activity_no": act_no,
                    "title": clean["title"],
                    "status": "updated"
                })
        else:
            new_act_id = uuid.uuid4()
            now_dt = datetime.utcnow()
            new_act = SubjectActivity(
                id=new_act_id,
                subject_id=subj_id,
                activity_no=act_no,
                title=clean["title"],
                unit_mapping=clean["unit_mapping"],
                estimated_time=clean["estimated_time"],
                mode=clean["mode"],
                file_naming=clean["file_naming"],
                why_this_activity=clean["why_this_activity"],
                instructions=clean["instructions"],
                ai_tools=clean["ai_tools"],
                learning_outcomes=clean["learning_outcomes"],
                submission_requirements=clean["submission_requirements"],
                rubric=clean["rubric"],
                is_released=True,
                released_at=now_dt,
                released_by="Auto Document Ingestion",
                created_at=now_dt,
                updated_at=now_dt
            )
            db.add(new_act)
            created_list.append({
                "id": str(new_act_id),
                "activity_no": act_no,
                "title": clean["title"],
                "unit_mapping": clean["unit_mapping"],
                "status": "created"
            })

    # 6. Commit primary database
    await db.commit()

    # 7. Sync new/updated activities to Neon PostgreSQL
    if created_list or updated_list:
        try:
            import asyncpg
            conn_neon = await asyncpg.connect(NEON_URI, ssl="require")
            for ca in created_list:
                act_row = await db.get(SubjectActivity, uuid.UUID(ca["id"]))
                if act_row:
                    await conn_neon.execute("""
                        INSERT INTO subject_activities (
                            id, subject_id, activity_no, title, unit_mapping, estimated_time,
                            mode, file_naming, why_this_activity, instructions, ai_tools,
                            learning_outcomes, submission_requirements, rubric, is_released,
                            released_at, released_by, created_at, updated_at
                        ) VALUES (
                            $1, $2, $3, $4, $5, $6, $7, $8, $9, $10, $11::json,
                            $12, $13, $14::json, $15, $16, $17, $18, $19
                        )
                        ON CONFLICT (id) DO UPDATE SET
                            title = EXCLUDED.title,
                            instructions = EXCLUDED.instructions,
                            rubric = EXCLUDED.rubric,
                            updated_at = EXCLUDED.updated_at
                    """,
                        act_row.id, act_row.subject_id, act_row.activity_no, act_row.title,
                        act_row.unit_mapping, act_row.estimated_time, act_row.mode,
                        act_row.file_naming, act_row.why_this_activity, act_row.instructions,
                        json.dumps(act_row.ai_tools), act_row.learning_outcomes,
                        act_row.submission_requirements, json.dumps(act_row.rubric),
                        act_row.is_released, act_row.released_at, act_row.released_by,
                        act_row.created_at, act_row.updated_at
                    )
            await conn_neon.close()
        except Exception as e:
            logger.warning(f"Could not mirror imported activities to Neon DB: {e}")

    return {
        "subject_id": str(subj_id),
        "subject_name": target_subject.name,
        "subject_code": target_subject.code,
        "total_activities_found": len(chunks),
        "created_count": len(created_list),
        "skipped_count": len(skipped_list),
        "updated_count": len(updated_list),
        "created_activities": created_list,
        "skipped_activities": skipped_list,
        "updated_activities": updated_list
    }
