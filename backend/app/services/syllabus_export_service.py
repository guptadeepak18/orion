import io
import re
from typing import Dict, Any, List
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

NAVY_HEX = "1B3A6B"
NAVY_COLOR = RGBColor(0x1B, 0x3A, 0x6B)
REPORTLAB_NAVY = colors.HexColor("#1B3A6B")

PROGRAM_OUTCOMES_REF = [
    {"po": "PO 1", "desc": "Apply knowledge of management theories and practices to solve business problems."},
    {"po": "PO 2", "desc": "Foster Analytical and critical thinking abilities for data-based decision making."},
    {"po": "PO 3", "desc": "Ability to develop Value based Leadership ability."},
    {"po": "PO 4", "desc": "Ability to understand, analyse and communicate global, economic, legal, and ethical aspects of business."},
    {"po": "PO 5", "desc": "Ability to lead themselves and others in the achievement of organizational goals, contributing effectively to a team environment."},
]


def set_cell_background(cell, hex_color=NAVY_HEX):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)


def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)


def add_docx_heading(doc: docx.Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.bold = True
    if level == 1:
        r.font.size = Pt(13)
        r.font.color.rgb = NAVY_COLOR
    else:
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def format_table_header(row, col_titles: List[str]):
    for i, cell in enumerate(row.cells):
        set_cell_background(cell, NAVY_HEX)
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.text = col_titles[i] if i < len(col_titles) else ""
        for r in p.runs:
            r.font.name = "Calibri"
            r.font.size = Pt(10)
            r.font.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def style_data_cell(cell, text: str, font_size: float = 9.5, bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT):
    set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
    p = cell.paragraphs[0]
    p.alignment = align
    p.text = text
    for r in p.runs:
        r.font.name = "Calibri"
        r.font.size = Pt(font_size)
        r.font.bold = bold


def generate_syllabus_docx(subject_name: str, subject_code: str, credits_info: str, syllabus_data: Dict[str, Any]) -> bytes:
    doc = docx.Document()

    # 1. Main Document Header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(f"{subject_name} ({subject_code}) — Course Syllabus")
    r.font.name = "Calibri"
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = NAVY_COLOR

    # Header Metadata Table
    meta_table = doc.add_table(rows=2, cols=3)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    format_table_header(meta_table.rows[0], ["Subject Name", "Subject Code", "Course Credits"])
    style_data_cell(meta_table.rows[1].cells[0], subject_name, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    style_data_cell(meta_table.rows[1].cells[1], subject_code, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    style_data_cell(meta_table.rows[1].cells[2], credits_info, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    # 2. Section 1: Course Overview
    add_docx_heading(doc, "Course Overview")
    overview = syllabus_data.get("course_overview", "")
    for para in overview.split("\n\n"):
        if para.strip():
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run(para.strip())
            r.font.name = "Calibri"
            r.font.size = Pt(10.5)

    # 3. Section 2: Course Outcomes & PO Matrix
    add_docx_heading(doc, "Course Outcomes & Program Outcomes Mapping Matrix")
    cos = syllabus_data.get("course_outcomes", [])
    if cos:
        co_table = doc.add_table(rows=len(cos) + 2, cols=8)
        co_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Title Header Row
        title_row = co_table.rows[0]
        title_row.cells[0].merge(title_row.cells[7])
        set_cell_background(title_row.cells[0], NAVY_HEX)
        p = title_row.cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("Mapping of Course Outcomes (COs) to Program Outcomes (POs)")
        r.font.name = "Calibri"
        r.font.size = Pt(10.5)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        # Column Header Row
        format_table_header(co_table.rows[1], ["Sr. No.", "Course Outcomes", "Bloom's Level", "PO1", "PO2", "PO3", "PO4", "PO5"])

        for idx, co in enumerate(cos):
            row = co_table.rows[idx + 2]
            style_data_cell(row.cells[0], str(idx + 1), align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
            style_data_cell(row.cells[1], co.get("statement", ""))
            style_data_cell(row.cells[2], co.get("blooms_level", "Knowledge"), align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
            style_data_cell(row.cells[3], str(co.get("po1", "-")), align=WD_ALIGN_PARAGRAPH.CENTER)
            style_data_cell(row.cells[4], str(co.get("po2", "-")), align=WD_ALIGN_PARAGRAPH.CENTER)
            style_data_cell(row.cells[5], str(co.get("po3", "-")), align=WD_ALIGN_PARAGRAPH.CENTER)
            style_data_cell(row.cells[6], str(co.get("po4", "-")), align=WD_ALIGN_PARAGRAPH.CENTER)
            style_data_cell(row.cells[7], str(co.get("po5", "-")), align=WD_ALIGN_PARAGRAPH.CENTER)

    # 4. Section 3: Topics / Units Breakdown
    add_docx_heading(doc, "Topics to be Covered (Course Content)")
    topics = syllabus_data.get("topics", [])
    if topics:
        u_table = doc.add_table(rows=len(topics) + 1, cols=3)
        u_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        format_table_header(u_table.rows[0], ["Unit No.", "Contents", "Session (Hours)"])
        for idx, unit in enumerate(topics):
            row = u_table.rows[idx + 1]
            style_data_cell(row.cells[0], str(unit.get("unit_number", idx + 1)), align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
            style_data_cell(row.cells[1], unit.get("content", ""))
            style_data_cell(row.cells[2], str(unit.get("hours", 6)), align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)

    # 5. Section 4: Industry Learning Activities (Case Studies & Live Assignments)
    add_docx_heading(doc, "Recommended Industry-Based Learning Activities")

    case_studies = syllabus_data.get("case_studies", [])
    if case_studies:
        add_docx_heading(doc, "Case Studies", level=2)
        cs_table = doc.add_table(rows=len(case_studies) + 1, cols=3)
        cs_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        format_table_header(cs_table.rows[0], ["Case Study Title", "Communication / Concept Covered", "Source / Access Link"])
        for idx, cs in enumerate(case_studies):
            row = cs_table.rows[idx + 1]
            style_data_cell(row.cells[0], cs.get("title", ""), bold=True)
            style_data_cell(row.cells[1], cs.get("concept", ""))
            style_data_cell(row.cells[2], cs.get("link", ""))

    live_assigns = syllabus_data.get("live_assignments", [])
    if live_assigns:
        add_docx_heading(doc, "Live Assignments", level=2)
        la_table = doc.add_table(rows=len(live_assigns) + 1, cols=3)
        la_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        format_table_header(la_table.rows[0], ["Assignment", "Task Description", "Source / Platform"])
        for idx, la in enumerate(live_assigns):
            row = la_table.rows[idx + 1]
            style_data_cell(row.cells[0], la.get("assignment", ""), bold=True)
            style_data_cell(row.cells[1], la.get("description", ""))
            style_data_cell(row.cells[2], la.get("platform", ""))

    software = syllabus_data.get("software_exposure", "")
    if software:
        add_docx_heading(doc, "Software Exposure", level=2)
        for line in software.split("\n"):
            if line.strip():
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(3)
                r = p.add_run(line.strip())
                r.font.name = "Calibri"
                r.font.size = Pt(10)

    # 6. Section 5 & 6: Scheme of Assessment & CCE Components
    add_docx_heading(doc, "Scheme of Assessment")
    ass_table = doc.add_table(rows=2, cols=3)
    ass_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    format_table_header(ass_table.rows[0], ["Comprehensive Concurrent Evaluation (CCE)", "Term End Examination (TEE)", "Total Marks"])
    style_data_cell(ass_table.rows[1].cells[0], f"{syllabus_data.get('cce_marks', 50)} Marks", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    style_data_cell(ass_table.rows[1].cells[1], f"{syllabus_data.get('tee_marks', 50)} Marks", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    style_data_cell(ass_table.rows[1].cells[2], f"{syllabus_data.get('total_marks', 100)} Marks", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)

    cce_comps = syllabus_data.get("cce_components", [])
    if cce_comps:
        add_docx_heading(doc, "Suggested CCE Components", level=2)
        for cce in cce_comps:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            r_bold = p.add_run(f"{cce.get('name', '')}: ")
            r_bold.font.name = "Calibri"
            r_bold.font.size = Pt(10)
            r_bold.font.bold = True

            r_desc = p.add_run(cce.get('details', ''))
            r_desc.font.name = "Calibri"
            r_desc.font.size = Pt(10)

    # 7. Section 7 & 8: Textbooks & Reference Books
    rec_books = syllabus_data.get("recommended_textbooks", [])
    if rec_books:
        add_docx_heading(doc, "Recommended Textbooks")
        for b in rec_books:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            citation = f"{b.get('author', '')} ({b.get('year', '2024')}). {b.get('name', '')}. {b.get('publisher', '')}."
            r = p.add_run(citation)
            r.font.name = "Calibri"
            r.font.size = Pt(10)

    ref_books = syllabus_data.get("reference_books", [])
    if ref_books:
        add_docx_heading(doc, "Reference Books")
        for b in ref_books:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            citation = f"{b.get('author', '')} ({b.get('year', '2024')}). {b.get('name', '')}. {b.get('publisher', '')}."
            r = p.add_run(citation)
            r.font.name = "Calibri"
            r.font.size = Pt(10)

    # 8. Section 9: Program Outcomes Reference Table
    add_docx_heading(doc, "Program Outcomes (PO Reference)")
    po_table = doc.add_table(rows=len(PROGRAM_OUTCOMES_REF) + 1, cols=2)
    po_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    format_table_header(po_table.rows[0], ["PO", "Program Outcome Description"])
    for idx, po in enumerate(PROGRAM_OUTCOMES_REF):
        row = po_table.rows[idx + 1]
        style_data_cell(row.cells[0], po["po"], align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
        style_data_cell(row.cells[1], po["desc"])

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def generate_syllabus_pdf(subject_name: str, subject_code: str, credits_info: str, syllabus_data: Dict[str, Any]) -> bytes:
    buffer = io.BytesIO()
    pdf = SimpleDocTemplate(buffer, pagesize=letter, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
    story = []

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        'TitleStyle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=16,
        leading=20,
        textColor=REPORTLAB_NAVY,
        spaceAfter=12,
    )

    h1_style = ParagraphStyle(
        'H1Style',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=12,
        leading=16,
        textColor=REPORTLAB_NAVY,
        spaceBefore=14,
        spaceAfter=6,
    )

    h2_style = ParagraphStyle(
        'H2Style',
        parent=styles['Heading3'],
        fontName='Helvetica-Bold',
        fontSize=10.5,
        leading=14,
        textColor=colors.HexColor("#333333"),
        spaceBefore=10,
        spaceAfter=4,
    )

    body_style = ParagraphStyle(
        'BodyTextCustom',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=13,
        spaceAfter=6,
    )

    tbl_hdr_style = ParagraphStyle(
        'TblHdr',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=9,
        leading=11,
        textColor=colors.white,
        alignment=1,  # Center
    )

    tbl_cell_style = ParagraphStyle(
        'TblCell',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9,
        leading=11,
    )

    tbl_cell_center = ParagraphStyle(
        'TblCellCenter',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9,
        leading=11,
        alignment=1,  # Center
    )

    # 1. Title
    story.append(Paragraph(f"{subject_name} ({subject_code}) — Course Syllabus", title_style))

    # Meta Table
    meta_data = [
        [Paragraph("Subject Name", tbl_hdr_style), Paragraph("Subject Code", tbl_hdr_style), Paragraph("Credits", tbl_hdr_style)],
        [Paragraph(subject_name, tbl_cell_center), Paragraph(subject_code, tbl_cell_center), Paragraph(credits_info, tbl_cell_center)]
    ]
    t_meta = Table(meta_data, colWidths=[240, 150, 150])
    t_meta.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), REPORTLAB_NAVY),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#D3D3D3")),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
    ]))
    story.append(t_meta)
    story.append(Spacer(1, 10))

    # 2. Overview
    story.append(Paragraph("Course Overview", h1_style))
    overview = syllabus_data.get("course_overview", "")
    for para in overview.split("\n\n"):
        if para.strip():
            story.append(Paragraph(para.strip(), body_style))

    # 3. Course Outcomes & PO Matrix
    cos = syllabus_data.get("course_outcomes", [])
    if cos:
        story.append(Paragraph("Course Outcomes & Program Outcomes Mapping Matrix", h1_style))
        co_rows = [
            [Paragraph("Sr.", tbl_hdr_style), Paragraph("Course Outcomes", tbl_hdr_style), Paragraph("Bloom's", tbl_hdr_style),
             Paragraph("PO1", tbl_hdr_style), Paragraph("PO2", tbl_hdr_style), Paragraph("PO3", tbl_hdr_style), Paragraph("PO4", tbl_hdr_style), Paragraph("PO5", tbl_hdr_style)]
        ]
        for idx, co in enumerate(cos):
            co_rows.append([
                Paragraph(str(idx + 1), tbl_cell_center),
                Paragraph(co.get("statement", ""), tbl_cell_style),
                Paragraph(co.get("blooms_level", "Knowledge"), tbl_cell_center),
                Paragraph(str(co.get("po1", "-")), tbl_cell_center),
                Paragraph(str(co.get("po2", "-")), tbl_cell_center),
                Paragraph(str(co.get("po3", "-")), tbl_cell_center),
                Paragraph(str(co.get("po4", "-")), tbl_cell_center),
                Paragraph(str(co.get("po5", "-")), tbl_cell_center),
            ])
        t_co = Table(co_rows, colWidths=[25, 235, 70, 35, 35, 35, 35, 35])
        t_co.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), REPORTLAB_NAVY),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#D3D3D3")),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('TOPPADDING', (0, 0), (-1, -1), 5),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
        ]))
        story.append(t_co)
        story.append(Spacer(1, 10))

    # 4. Units / Topics
    topics = syllabus_data.get("topics", [])
    if topics:
        story.append(Paragraph("Topics to be Covered (Course Content)", h1_style))
        u_rows = [
            [Paragraph("Unit No.", tbl_hdr_style), Paragraph("Contents", tbl_hdr_style), Paragraph("Session (Hours)", tbl_hdr_style)]
        ]
        for idx, unit in enumerate(topics):
            u_rows.append([
                Paragraph(str(unit.get("unit_number", idx + 1)), tbl_cell_center),
                Paragraph(unit.get("content", ""), tbl_cell_style),
                Paragraph(str(unit.get("hours", 6)), tbl_cell_center)
            ])
        t_u = Table(u_rows, colWidths=[60, 400, 80])
        t_u.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), REPORTLAB_NAVY),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#D3D3D3")),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('TOPPADDING', (0, 0), (-1, -1), 5),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
        ]))
        story.append(t_u)
        story.append(Spacer(1, 10))

    # 5. Industry Activities
    case_studies = syllabus_data.get("case_studies", [])
    live_assigns = syllabus_data.get("live_assignments", [])
    if case_studies or live_assigns:
        story.append(Paragraph("Recommended Industry-Based Learning Activities", h1_style))
        if case_studies:
            story.append(Paragraph("Case Studies", h2_style))
            cs_rows = [[Paragraph("Case Study Title", tbl_hdr_style), Paragraph("Concept Covered", tbl_hdr_style), Paragraph("Source / Link", tbl_hdr_style)]]
            for cs in case_studies:
                cs_rows.append([Paragraph(cs.get("title", ""), tbl_cell_style), Paragraph(cs.get("concept", ""), tbl_cell_style), Paragraph(cs.get("link", ""), tbl_cell_style)])
            t_cs = Table(cs_rows, colWidths=[180, 200, 160])
            t_cs.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), REPORTLAB_NAVY),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#D3D3D3")),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('TOPPADDING', (0, 0), (-1, -1), 4),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ]))
            story.append(t_cs)

    # 6. Scheme of Assessment
    story.append(Paragraph("Scheme of Assessment", h1_style))
    ass_rows = [
        [Paragraph("CCE Marks", tbl_hdr_style), Paragraph("TEE Marks", tbl_hdr_style), Paragraph("Total Marks", tbl_hdr_style)],
        [Paragraph(f"{syllabus_data.get('cce_marks', 50)} Marks", tbl_cell_center), Paragraph(f"{syllabus_data.get('tee_marks', 50)} Marks", tbl_cell_center), Paragraph(f"{syllabus_data.get('total_marks', 100)} Marks", tbl_cell_center)]
    ]
    t_ass = Table(ass_rows, colWidths=[180, 180, 180])
    t_ass.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), REPORTLAB_NAVY),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#D3D3D3")),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('TOPPADDING', (0, 0), (-1, -1), 5),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
    ]))
    story.append(t_ass)

    cce_comps = syllabus_data.get("cce_components", [])
    if cce_comps:
        story.append(Paragraph("Suggested CCE Components", h2_style))
        for cce in cce_comps:
            story.append(Paragraph(f"<b>{cce.get('name', '')}:</b> {cce.get('details', '')}", body_style))

    # 7 & 8 Textbooks & Reference Books
    rec_books = syllabus_data.get("recommended_textbooks", [])
    if rec_books:
        story.append(Paragraph("Recommended Textbooks", h1_style))
        for b in rec_books:
            story.append(Paragraph(f"• {b.get('author', '')} ({b.get('year', '2024')}). <i>{b.get('name', '')}</i>. {b.get('publisher', '')}.", body_style))

    # 9. Program Outcomes Reference Table
    story.append(Paragraph("Program Outcomes (PO Reference)", h1_style))
    po_rows = [[Paragraph("PO", tbl_hdr_style), Paragraph("Program Outcome Description", tbl_hdr_style)]]
    for po in PROGRAM_OUTCOMES_REF:
        po_rows.append([Paragraph(po["po"], tbl_cell_center), Paragraph(po["desc"], tbl_cell_style)])
    t_po = Table(po_rows, colWidths=[60, 480])
    t_po.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), REPORTLAB_NAVY),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#D3D3D3")),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
    ]))
    story.append(t_po)

    pdf.build(story)
    return buffer.getvalue()
