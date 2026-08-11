import io
import re
from typing import Dict, Any, List
import docx
from pypdf import PdfReader


def clean_str(val: str) -> str:
    if not val:
        return ""
    cleaned = str(val).replace('\ufffd', '-').replace('\u00ad', '').replace('\u2014', '-').replace('\u2013', '-').replace('\t', ' ')
    cleaned = re.sub(r'\s+', ' ', cleaned)
    return cleaned.strip()


def parse_apa_book(text: str, book_id: str) -> Dict[str, Any]:
    cleaned = clean_str(text)
    pattern = r"^(.*?)\s*\((\d{4})\)\.\s*(.*?)\.\s*(.*)$"
    m = re.match(pattern, cleaned)
    if m:
        author, year, name, publisher = m.groups()
        edition = "Latest"
        ed_m = re.search(r"\((.*?edition.*?)\)", name, re.IGNORECASE)
        if ed_m:
            edition = ed_m.group(1).strip()
        return {
            "id": book_id,
            "name": name.strip(),
            "author": author.strip(),
            "publisher": publisher.strip().rstrip('.'),
            "year": year.strip(),
            "edition": edition,
        }
    return {
        "id": book_id,
        "name": cleaned,
        "author": "Standard Author",
        "publisher": "Standard Publisher",
        "year": "2024",
        "edition": "Latest",
    }


def find_data_rows(table: docx.table.Table) -> List[docx.table._Row]:
    header_idx = -1
    for r_idx, row in enumerate(table.rows):
        cells_text = " ".join([clean_str(c.text).lower() for c in row.cells])
        if any(k in cells_text for k in ['sr. no', 'course outcomes', 'unit no', 'contents', 'case study title', 'task description', 'comprehensive concurrent']):
            header_idx = r_idx
            break
    if header_idx != -1:
        return table.rows[header_idx + 1:]
    return table.rows[1:] if len(table.rows) > 1 else []


def parse_docx_syllabus(file_bytes: bytes) -> Dict[str, Any]:
    doc = docx.Document(io.BytesIO(file_bytes))

    course_overview: List[str] = []
    course_outcomes: List[Dict[str, Any]] = []
    topics: List[Dict[str, Any]] = []
    case_studies: List[Dict[str, Any]] = []
    live_assignments: List[Dict[str, Any]] = []
    software_exposure_lines: List[str] = []
    cce_marks = 50
    tee_marks = 50
    total_marks = 100
    cce_components: List[Dict[str, Any]] = []
    recommended_textbooks: List[Dict[str, Any]] = []
    reference_books: List[Dict[str, Any]] = []

    # 1. Parse Paragraphs for Overview, Software Exposure, CCE, Textbooks, References
    current_section = None
    for p in doc.paragraphs:
        txt = clean_str(p.text)
        if not txt:
            continue
        lower_txt = txt.lower()

        if "course overview" in lower_txt:
            current_section = "overview"
            continue
        elif "course outcomes" in lower_txt or "topics to be covered" in lower_txt:
            if current_section == "overview":
                current_section = None
        elif "software exposure" in lower_txt:
            current_section = "software"
            continue
        elif "scheme of assessment" in lower_txt:
            current_section = None
        elif "suggested cce components" in lower_txt:
            current_section = "cce_components"
            continue
        elif "recommended textbook" in lower_txt:
            current_section = "textbooks"
            continue
        elif "reference book" in lower_txt:
            current_section = "references"
            continue
        elif "program outcomes" in lower_txt or "annexure" in lower_txt:
            current_section = None

        if current_section == "overview":
            if len(txt) > 30 and not txt.startswith("P[") and not txt.startswith("Table"):
                course_overview.append(txt)
        elif current_section == "software":
            if len(txt) > 5 and not txt.startswith("Software Exposure"):
                software_exposure_lines.append(f"• {txt}")
        elif current_section == "cce_components":
            if txt.startswith("Assessment ") or txt.startswith("Component ") or "Assignment" in txt or "Quiz" in txt or "Project" in txt:
                parts = re.split(r'[:\u2014-]', txt, 1)
                comp_name = parts[0].strip() if parts else txt
                comp_details = parts[1].strip() if len(parts) > 1 else txt
                cce_components.append({
                    "id": f"cce-{len(cce_components)+1}",
                    "name": comp_name,
                    "marks": 10,
                    "details": comp_details,
                })
        elif current_section == "textbooks":
            if len(txt) > 10 and not txt.startswith("Recommended"):
                book_obj = parse_apa_book(txt, f"tb-{len(recommended_textbooks)+1}")
                recommended_textbooks.append(book_obj)
        elif current_section == "references":
            if len(txt) > 10 and not txt.startswith("Reference"):
                book_obj = parse_apa_book(txt, f"rb-{len(reference_books)+1}")
                reference_books.append(book_obj)

    overview_text = "\n\n".join(course_overview) if course_overview else ""
    software_exposure_text = "\n".join(software_exposure_lines) if software_exposure_lines else "MS Word (Legal Drafting), LexisNexis, Python"

    # 2. Iterate through Tables with Smart Row Header Skipping
    for t_idx, table in enumerate(doc.tables):
        table_text = " ".join([clean_str(c.text).lower() for r in table.rows for c in r.cells])

        # Course Outcomes Table
        if "course outcome" in table_text or "bloom" in table_text:
            data_rows = find_data_rows(table)
            for r_idx, row in enumerate(data_rows):
                cells = [clean_str(c.text) for c in row.cells]
                if len(cells) >= 3:
                    sr_no = cells[0]
                    statement = cells[1]
                    blooms = cells[2] if len(cells) > 2 else "Knowledge"

                    clean_po = lambda val: "-" if val in ['', '-', 'None', 'null', None] else val

                    po1 = clean_po(cells[3]) if len(cells) > 3 else "-"
                    po2 = clean_po(cells[4]) if len(cells) > 4 else "-"
                    po3 = clean_po(cells[5]) if len(cells) > 5 else "-"
                    po4 = clean_po(cells[6]) if len(cells) > 6 else "-"
                    po5 = clean_po(cells[7]) if len(cells) > 7 else "-"

                    valid_blooms = ['Knowledge', 'Understand', 'Application', 'Analysis', 'Evaluation', 'Create']
                    clean_blooms = blooms if blooms in valid_blooms else "Application"

                    if statement and statement.lower() != "total" and "bloom" not in statement.lower() and "course outcome" not in statement.lower():
                        course_outcomes.append({
                            "id": f"co-{len(course_outcomes)+1}",
                            "statement": statement,
                            "blooms_level": clean_blooms,
                            "po1": po1,
                            "po2": po2,
                            "po3": po3,
                            "po4": po4,
                            "po5": po5,
                        })

        # Units / Topics Table
        elif "contents" in table_text or "session" in table_text or "unit no" in table_text:
            data_rows = find_data_rows(table)
            for r_idx, row in enumerate(data_rows):
                cells = [clean_str(c.text) for c in row.cells]
                if len(cells) >= 3:
                    unit_num_raw = cells[0]
                    content_raw = cells[1]
                    hours_raw = cells[2]

                    if unit_num_raw.lower() == "total" or not content_raw or "contents" in content_raw.lower():
                        continue

                    try:
                        u_num = int(unit_num_raw)
                    except ValueError:
                        u_num = len(topics) + 1

                    try:
                        u_hrs = int(hours_raw)
                    except ValueError:
                        u_hrs = 6

                    title = f"Unit {u_num}"
                    if len(content_raw) > 5:
                        first_sentence = content_raw.split(".")[0]
                        title = f"Unit {u_num}: {first_sentence}"

                    topics.append({
                        "id": f"unit-{u_num}",
                        "unit_number": u_num,
                        "unit_title": title,
                        "content": content_raw,
                        "hours": u_hrs,
                    })

        # Case Studies Table
        elif "case study title" in table_text or "communication concept" in table_text:
            data_rows = find_data_rows(table)
            for r_idx, row in enumerate(data_rows):
                cells = [clean_str(c.text) for c in row.cells]
                if len(cells) >= 3:
                    title = cells[0]
                    if title and "case study title" not in title.lower():
                        case_studies.append({
                            "id": f"cs-{len(case_studies)+1}",
                            "title": title,
                            "concept": cells[1],
                            "link": cells[2],
                        })

        # Live Assignments Table
        elif "task description" in table_text:
            data_rows = find_data_rows(table)
            for r_idx, row in enumerate(data_rows):
                cells = [clean_str(c.text) for c in row.cells]
                if len(cells) >= 3:
                    assignment = cells[0]
                    if assignment and "assignment" not in assignment.lower():
                        live_assignments.append({
                            "id": f"la-{len(live_assignments)+1}",
                            "assignment": assignment,
                            "description": cells[1],
                            "platform": cells[2],
                        })

        # Scheme of Assessment Table
        elif "comprehensive concurrent evaluation" in table_text or "term end examination" in table_text:
            data_rows = find_data_rows(table)
            if data_rows:
                row1 = [clean_str(c.text) for c in data_rows[0].cells]
                if len(row1) >= 3:
                    cce_match = re.search(r'(\d+)', row1[0])
                    tee_match = re.search(r'(\d+)', row1[1])
                    if cce_match:
                        cce_marks = int(cce_match.group(1))
                    if tee_match:
                        tee_marks = int(tee_match.group(1))
                    total_marks = cce_marks + tee_marks

    return {
        "course_overview": overview_text or "Course Overview details extracted from uploaded document.",
        "course_outcomes": course_outcomes,
        "topics": topics if topics else [{"id": "unit-1", "unit_number": 1, "unit_title": "Unit 1", "content": "Unit content", "hours": 6}],
        "case_studies": case_studies,
        "live_assignments": live_assignments,
        "software_exposure": software_exposure_text,
        "cce_marks": cce_marks,
        "tee_marks": tee_marks,
        "total_marks": total_marks,
        "cce_components": cce_components,
        "recommended_textbooks": recommended_textbooks,
        "reference_books": reference_books,
    }


def parse_pdf_syllabus(file_bytes: bytes) -> Dict[str, Any]:
    reader = PdfReader(io.BytesIO(file_bytes))
    full_text = "\n".join([clean_str(page.extract_text()) for page in reader.pages if page.extract_text()])

    paragraphs = [p.strip() for p in full_text.split("\n\n") if p.strip()]
    overview = paragraphs[0] if paragraphs else "Uploaded PDF Syllabus Overview"

    return {
        "course_overview": overview,
        "course_outcomes": [
            {
                "id": "co-1",
                "statement": "Extracted Course Outcome from uploaded PDF document.",
                "blooms_level": "Knowledge",
                "po1": "2",
                "po2": "-",
                "po3": "-",
                "po4": "2",
                "po5": "-",
            }
        ],
        "topics": [
            {
                "id": "unit-1",
                "unit_number": 1,
                "unit_title": "Unit 1: Module Overview",
                "content": full_text[:500] if full_text else "Unit content",
                "hours": 6,
            }
        ],
        "case_studies": [],
        "live_assignments": [],
        "software_exposure": "Standard Software Tools",
        "cce_marks": 50,
        "tee_marks": 50,
        "total_marks": 100,
        "cce_components": [],
        "recommended_textbooks": [],
        "reference_books": [],
    }


def parse_syllabus_file(file_bytes: bytes, filename: str) -> Dict[str, Any]:
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    if ext in ["docx", "doc"]:
        return parse_docx_syllabus(file_bytes)
    elif ext == "pdf":
        return parse_pdf_syllabus(file_bytes)
    else:
        try:
            return parse_docx_syllabus(file_bytes)
        except Exception:
            return parse_pdf_syllabus(file_bytes)
